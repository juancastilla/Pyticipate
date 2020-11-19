import core.scenario.scenario_base as sb

import base64
import datetime as dt
import flopy
import geopandas as gpd
import numpy as np
import os
import pandas as pd
import pickle

from math import pi

import matplotlib as mpl
mpl.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
from matplotlib.gridspec import GridSpec
import matplotlib.ticker as ticker

pd.plotting.register_matplotlib_converters()

import seaborn as sns

import docx
from docxcompose.composer import Composer
from docx import Document as Document_compose
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.section import WD_SECTION
from docx.shared import Cm
from docx.shared import Pt


from multiprocessing import Pool

import sys
sys.path.append(os.path.dirname(os.path.realpath(__file__)))

import model.process.threadedfunctions as t_funcs

from model.surface_water import dam_operations as dam_ops

class Scenario(sb.ScenarioBase):

    def dispose(self):
        pass

    def path_for_resource(self, resource):
        return os.path.join(self.static_dir, resource)

    def init_file_system(self, working_dir):
        scenario_dir = os.path.abspath(os.path.dirname(__file__))
        self.static_dir = os.path.join(scenario_dir, 'static')
        self.layer_dir = os.path.join(scenario_dir, 'layers')
        self.model_data_dir = os.path.join(scenario_dir, 'model', 'data')
        self.model_geodata_dir = os.path.join(scenario_dir, 'model', 'geodata')
        self.model_modflow_bin_dir = os.path.join(scenario_dir, 'model', 'modflow', 'bin')
        self.model_modflow_data_dir = os.path.join(scenario_dir, 'model', 'modflow', 'data')
        self.model_report_data_dir = os.path.join(self.model_data_dir, 'report')
        self.working_dir = working_dir
        self.model_output_dir = os.path.join(working_dir, 'output')
        os.makedirs(self.model_output_dir)
        self.model_report_dir = os.path.join(working_dir, 'report')
        os.makedirs(self.model_report_dir)
        self.model_modflow_load_dir = os.path.join(working_dir, 'modflow', 'load')
        os.makedirs(self.model_modflow_load_dir)
        self.model_modflow_run_dir = os.path.join(working_dir, 'modflow', 'run')
        os.makedirs(self.model_modflow_run_dir)

    def __init__(self, working_dir):
        self.report = None
        # Initialize scenario file system
        self.init_file_system(working_dir)
        # Define temporal characteristics
        start = dt.datetime(2019,9,30)
        end = dt.datetime(2044,9,30)
        timestep = sb.TimeStep.Parse('Annual')
        multi_step = True
        super().__init__(start, end, timestep, multi_step)
        self.set_run_status(0.0, 'Inicializando...')

        # Define Input Groups

        input_group_0_id = 'input_group_0'
        input_group_0a_id = 'input_group_0_subgroup_a'
        super().add_group(sb.GroupDef(input_group_0_id,'Series Hidrológicas',None),True)
        super().add_group(sb.GroupDef(input_group_0a_id,'Series Hidrológicas',input_group_0_id,os.path.join(self.static_dir,'html/input/group_0.html')),True)
        input_group_1_id = 'input_group_1'
        input_group_1a_id = 'input_group_1_subgroup_a'
        input_group_1b_id = 'input_group_1_subgroup_b'
        input_group_1c_id = 'input_group_1_subgroup_c'
        super().add_group(sb.GroupDef(input_group_1_id,'Swap Hídrico'),True)
        super().add_group(sb.GroupDef(input_group_1a_id,'Candelaria 175 L/s a Aguas Chañar?',input_group_1_id,os.path.join(self.static_dir,'html/input/group_1a.html')),True)
        super().add_group(sb.GroupDef(input_group_1b_id,'Caserones 200 L/s a Rio Ramadilla?',input_group_1_id,os.path.join(self.static_dir,'html/input/group_1b.html')),True)
        super().add_group(sb.GroupDef(input_group_1c_id,'Redistribuir Agua Superficial Distritos 8/9?',input_group_1_id,os.path.join(self.static_dir,'html/input/group_1c.html')),True)
        input_group_2_id = 'input_group_2'
        input_group_2a_id = 'input_group_2_subgroup_a'
        input_group_2b_id = 'input_group_2_subgroup_b'
        input_group_2c_id = 'input_group_2_subgroup_c'
        super().add_group(sb.GroupDef(input_group_2_id,'Inversion infraestructura'),True)
        super().add_group(sb.GroupDef(input_group_2a_id,'Lautaro 2.0?',input_group_2_id,os.path.join(self.static_dir,'html/input/group_2a.html')),True)
        super().add_group(sb.GroupDef(input_group_2b_id,'Entubamiento canales de riego?',input_group_2_id,os.path.join(self.static_dir,'html/input/group_2b.html')),True)
        super().add_group(sb.GroupDef(input_group_2c_id,'Operacion Desaladora (Año desactivación)',input_group_2_id,os.path.join(self.static_dir,'html/input/group_2c.html')),True)
        input_group_3_id = 'input_group_3'
        input_group_3a_id = 'input_group_3_subgroup_a'
        super().add_group(sb.GroupDef(input_group_3_id,'Recarga Artificial Acuífero'),True)
        super().add_group(sb.GroupDef(input_group_3a_id,'Obras de recarga artificial en cauce Rio Copiapó?',input_group_3_id,os.path.join(self.static_dir,'html/input/group_3.html')),True)
        input_group_4_id = 'input_group_4'
        input_group_4a_id = 'input_group_4_subgroup_a'
        super().add_group(sb.GroupDef(input_group_4_id,'Prorrata aguas subterráneas'),True)
        super().add_group(sb.GroupDef(input_group_4a_id,'Prorrateo aguas subterráneas minería y agricultura?',input_group_4_id,os.path.join(self.static_dir,'html/input/group_4.html')),True)
        input_group_5_id = 'input_group_5'
        input_group_5a_id = 'input_group_5_subgroup_a'
        super().add_group(sb.GroupDef(input_group_5_id,'Reutilización aguas grises domésticas'),True)
        super().add_group(sb.GroupDef(input_group_5a_id,'Incentivos o infraestructura para reutilizar aguas grises?',input_group_5_id,os.path.join(self.static_dir,'html/input/group_5.html')),True)

        # Define Output Groups

        output_group_1_id = 'output_group_1'
        super().add_group(sb.GroupDef(output_group_1_id,'Cambio en volumen embalsado'),False)
        output_group_2_id = 'output_group_2'
        super().add_group(sb.GroupDef(output_group_2_id,'Profundidad napa y costo de bombeo'),False)
        output_group_3_id = 'output_group_3'
        super().add_group(sb.GroupDef(output_group_3_id,'Recargas riego'),False)
        output_group_4_id = 'output_group_4'
        super().add_group(sb.GroupDef(output_group_4_id,'Recargas Rio Copiapo'),False)
        output_group_5_id = 'output_group_5'
        super().add_group(sb.GroupDef(output_group_5_id,'Satisfaccion D canales'),False)
        output_group_6_id = 'output_group_6'
        super().add_group(sb.GroupDef(output_group_6_id,'Satisfaccion D mixta'),False)
        output_group_7_id = 'output_group_7'
        super().add_group(sb.GroupDef(output_group_7_id,'Desaladora estatal'),False)

        # Define Inputs

        super().add_input(sb.InputSingleSelectionDef('input_0','Series',input_group_0a_id,['Historico','Reduccion 50%','Invertida'],'Historico'))

        super().add_input(sb.InputBooleanDef('input_1a01','Swap Candelaria-AChañar',input_group_1a_id,False))
        super().add_input(sb.InputSingleSelectionDef('input_1a02','Dejar de bombear pozos',input_group_1a_id,['PC','PU'],'PC'))
        super().add_input(sb.InputBooleanDef('input_1b01','Swap Caserones-Ramadilla',input_group_1b_id,False))
        super().add_input(sb.InputBooleanDef('input_1c01','Redistribucion oferta D89?',input_group_1c_id,False))
        super().add_input(sb.InputSingleSelectionDef('input_1c02','Oferta D89',input_group_1c_id,['A D17','A Aguas Chañar','A S5 y S6 por rio','A S5 por tuberia (excedente a rio)','A S5 por tuberia (excedente a recarga artificial)'],'A D17'))

        super().add_input(sb.InputBooleanDef('input_2a01','Lautaro 2.0',input_group_2a_id,False))
        super().add_input(sb.InputBooleanDef('input_2b01','Entubamiento canales distritos',input_group_2b_id,False))
        super().add_input(sb.InputBoundNumericDef('input_2c01','Pozo PC01',input_group_2c_id,2019,2044,2024))
        super().add_input(sb.InputBoundNumericDef('input_2c02','Pozo PC02',input_group_2c_id,2019,2044,2029))
        super().add_input(sb.InputBoundNumericDef('input_2c03','Pozo PC03',input_group_2c_id,2019,2044,2029))
        super().add_input(sb.InputBoundNumericDef('input_2c04','Pozo PC04',input_group_2c_id,2019,2044,2029))
        super().add_input(sb.InputBoundNumericDef('input_2c05','Pozo PC05',input_group_2c_id,2019,2044,2029))
        super().add_input(sb.InputBoundNumericDef('input_2c06','Pozo PC06',input_group_2c_id,2019,2044,2029))
        super().add_input(sb.InputBoundNumericDef('input_2c07','Pozo PC07',input_group_2c_id,2019,2044,2030))
        super().add_input(sb.InputBoundNumericDef('input_2c08','Pozo PU01',input_group_2c_id,2019,2044,2024))
        super().add_input(sb.InputBoundNumericDef('input_2c09','Pozo PU02',input_group_2c_id,2019,2044,2024))
        super().add_input(sb.InputBoundNumericDef('input_2c10','Pozo PU03',input_group_2c_id,2019,2044,2024))
        super().add_input(sb.InputBoundNumericDef('input_2c11','Pozo CE01',input_group_2c_id,2019,2044,2024))

        super().add_input(sb.InputBooleanDef('input_3a01','Recarga artificial S3 Nantoco',input_group_3a_id,False))
        super().add_input(sb.InputBooleanDef('input_3a02','Recarga artificial S4 (antes Kaukari)',input_group_3a_id,False))
        super().add_input(sb.InputBooleanDef('input_3a03','Recarga artificial S4 (despues Kaukari)',input_group_3a_id,False))
        super().add_input(sb.InputBooleanDef('input_3a04','Recarga artificial S5 Piedra Colgada',input_group_3a_id,False))
        super().add_input(sb.InputNumericDef('input_3a05','Tasa recarga nominal (m/d)',input_group_3a_id,1,0.1,allowNegative=False))

        super().add_input(sb.InputBoundNumericDef('input_4a01','Prorrateo S2 agricultura',input_group_4a_id,0,1,1.00,0.01))
        super().add_input(sb.InputBoundNumericDef('input_4a02','Prorrateo S3 agricultura',input_group_4a_id,0,1,0.60,0.01))
        super().add_input(sb.InputBoundNumericDef('input_4a03','Prorrateo S4 agricultura',input_group_4a_id,0,1,0.60,0.01))
        super().add_input(sb.InputBoundNumericDef('input_4a04','Prorrateo S5 agricultura',input_group_4a_id,0,1,0.60,0.01))
        super().add_input(sb.InputBoundNumericDef('input_4a05','Prorrateo S6 agricultura',input_group_4a_id,0,1,1.00,0.01))
        super().add_input(sb.InputBoundNumericDef('input_4a06','Prorrateo S2 mineria',input_group_4a_id,0,1,1.00,0.01))
        super().add_input(sb.InputBoundNumericDef('input_4a07','Prorrateo S3 mineria',input_group_4a_id,0,1,1.00,0.01))
        super().add_input(sb.InputBoundNumericDef('input_4a08','Prorrateo S4 mineria',input_group_4a_id,0,1,1.00,0.01))
        super().add_input(sb.InputBoundNumericDef('input_4a09','Prorrateo S5 mineria',input_group_4a_id,0,1,1.00,0.01))
        super().add_input(sb.InputBoundNumericDef('input_4a10','Prorrateo S6 mineria',input_group_4a_id,0,1,1.00,0.01))
        super().add_input(sb.InputBoundNumericDef('input_4a11','Prorrateo monitoreo',input_group_4a_id,0.1,0.9,0.9,0.1))
        super().add_input(sb.InputBoundNumericDef('input_4a12','Prorrateo multas',input_group_4a_id,0.1,0.9,0.9,0.1))

        super().add_input(sb.InputBoundNumericDef('input_5a01','% reduccion consumo AP',input_group_5a_id,0,100,0,1))

        # Define Outputs

        output_decimal_places = 3

        sector_names = ['Sector 2', 'Sector 3', 'Sector 4', 'Sector 5', 'Sector 6']
        super().add_output(sb.OutputHistogramDef('output_02','Cambio volumen en almacenamiento Asub año 25',output_group_1_id,sector_names,sb.ChartDisplay('Sector acuífero','Cambio (M m3)',output_decimal_places)))
        super().add_output(sb.OutputHistogramDef('output_03','Cambio volumen (Embalse Lautaro)',output_group_1_id,sector_names,sb.ChartDisplay('Sector acuífero','Cambio (factor volumen embalse)',output_decimal_places)))

        super().add_output(sb.OutputHistogramDef('output_04','Cambio en profundidad napa freática',output_group_2_id,sector_names,sb.ChartDisplay('Sector acuífero','Cambio promedio (m)',output_decimal_places)))
        super().add_output(sb.OutputHistogramDef('output_05','Costos promedios de bombeo',output_group_2_id,sector_names,sb.ChartDisplay('Sector acuífero','Costo promedio (CLP/K m3)',output_decimal_places)))

        temporal_units_label = 'Fecha'
        flow_units_label = 'Litros/segundo'
        percentage_units_label = '%'

        super().add_output(sb.OutputTimeSeriesDef('output_10','Recarga desde riego (Sector 2)',output_group_3_id,sb.ChartDisplay(temporal_units_label,flow_units_label,output_decimal_places)))
        super().add_output(sb.OutputTimeSeriesDef('output_11','Recarga desde riego (Sector 3)',output_group_3_id,sb.ChartDisplay(temporal_units_label,flow_units_label,output_decimal_places)))
        super().add_output(sb.OutputTimeSeriesDef('output_12','Recarga desde riego (Sector 4)',output_group_3_id,sb.ChartDisplay(temporal_units_label,flow_units_label,output_decimal_places)))
        super().add_output(sb.OutputTimeSeriesDef('output_13','Recarga desde riego (Sector 5)',output_group_3_id,sb.ChartDisplay(temporal_units_label,flow_units_label,output_decimal_places)))
        super().add_output(sb.OutputTimeSeriesDef('output_14','Recarga desde riego (Sector 6)',output_group_3_id,sb.ChartDisplay(temporal_units_label,flow_units_label,output_decimal_places)))

        super().add_output(sb.OutputTimeSeriesDef('output_20','Recarga desde el rio Copiapó (Sector 2)',output_group_4_id,sb.ChartDisplay(temporal_units_label,flow_units_label,output_decimal_places)))
        super().add_output(sb.OutputTimeSeriesDef('output_21','Recarga desde el rio Copiapó (Sector 3)',output_group_4_id,sb.ChartDisplay(temporal_units_label,flow_units_label,output_decimal_places)))
        super().add_output(sb.OutputTimeSeriesDef('output_22','Recarga desde el rio Copiapó (Sector 4)',output_group_4_id,sb.ChartDisplay(temporal_units_label,flow_units_label,output_decimal_places)))
        super().add_output(sb.OutputTimeSeriesDef('output_23','Recarga desde el rio Copiapó (Sector 5)',output_group_4_id,sb.ChartDisplay(temporal_units_label,flow_units_label,output_decimal_places)))
        super().add_output(sb.OutputTimeSeriesDef('output_24','Recarga desde el rio Copiapó (Sector 6)',output_group_4_id,sb.ChartDisplay(temporal_units_label,flow_units_label,output_decimal_places)))

        super().add_output(sb.OutputTimeSeriesDef('output_30','Demanda satisfecha ASup distrito de riego (Canales D1)',output_group_5_id,sb.ChartDisplay(temporal_units_label,percentage_units_label,output_decimal_places)))
        super().add_output(sb.OutputTimeSeriesDef('output_31','Demanda satisfecha ASup distrito de riego (Canales D2)',output_group_5_id,sb.ChartDisplay(temporal_units_label,percentage_units_label,output_decimal_places)))
        super().add_output(sb.OutputTimeSeriesDef('output_32','Demanda satisfecha ASup distrito de riego (Canales D3)',output_group_5_id,sb.ChartDisplay(temporal_units_label,percentage_units_label,output_decimal_places)))
        super().add_output(sb.OutputTimeSeriesDef('output_33','Demanda satisfecha ASup distrito de riego (Canales D4)',output_group_5_id,sb.ChartDisplay(temporal_units_label,percentage_units_label,output_decimal_places)))
        super().add_output(sb.OutputTimeSeriesDef('output_34','Demanda satisfecha ASup distrito de riego (Canales D5)',output_group_5_id,sb.ChartDisplay(temporal_units_label,percentage_units_label,output_decimal_places)))
        super().add_output(sb.OutputTimeSeriesDef('output_35','Demanda satisfecha ASup distrito de riego (Canales D6)',output_group_5_id,sb.ChartDisplay(temporal_units_label,percentage_units_label,output_decimal_places)))
        super().add_output(sb.OutputTimeSeriesDef('output_36','Demanda satisfecha ASup distrito de riego (Canales D7)',output_group_5_id,sb.ChartDisplay(temporal_units_label,percentage_units_label,output_decimal_places)))
        super().add_output(sb.OutputTimeSeriesDef('output_37','Demanda satisfecha ASup distrito de riego (Canales D89)',output_group_5_id,sb.ChartDisplay(temporal_units_label,percentage_units_label,output_decimal_places)))

        super().add_output(sb.OutputTimeSeriesDef('output_40','Demanda satisfecha ASup distrito de riego (Mixto D1)',output_group_6_id,sb.ChartDisplay(temporal_units_label,percentage_units_label,output_decimal_places)))
        super().add_output(sb.OutputTimeSeriesDef('output_41','Demanda satisfecha ASup distrito de riego (Mixto D2)',output_group_6_id,sb.ChartDisplay(temporal_units_label,percentage_units_label,output_decimal_places)))
        super().add_output(sb.OutputTimeSeriesDef('output_42','Demanda satisfecha ASup distrito de riego (Mixto D3)',output_group_6_id,sb.ChartDisplay(temporal_units_label,percentage_units_label,output_decimal_places)))
        super().add_output(sb.OutputTimeSeriesDef('output_43','Demanda satisfecha ASup distrito de riego (Mixto D4)',output_group_6_id,sb.ChartDisplay(temporal_units_label,percentage_units_label,output_decimal_places)))
        super().add_output(sb.OutputTimeSeriesDef('output_44','Demanda satisfecha ASup distrito de riego (Mixto D5)',output_group_6_id,sb.ChartDisplay(temporal_units_label,percentage_units_label,output_decimal_places)))
        super().add_output(sb.OutputTimeSeriesDef('output_45','Demanda satisfecha ASup distrito de riego (Mixto D6)',output_group_6_id,sb.ChartDisplay(temporal_units_label,percentage_units_label,output_decimal_places)))
        super().add_output(sb.OutputTimeSeriesDef('output_46','Demanda satisfecha ASup distrito de riego (Mixto D7)',output_group_6_id,sb.ChartDisplay(temporal_units_label,percentage_units_label,output_decimal_places)))
        super().add_output(sb.OutputTimeSeriesDef('output_47','Demanda satisfecha ASup distrito de riego (Mixto D89)',output_group_6_id,sb.ChartDisplay(temporal_units_label,percentage_units_label,output_decimal_places)))

        super().add_output(sb.OutputHistogramDef('output_50','Operacion planta desaladora estatal',output_group_7_id,[str(y) for y in range(0, 25)],sb.ChartDisplay('Años',flow_units_label,output_decimal_places)))

        # Define Layers

        # Layer 01 Groundwater Heads (GeoTIFF)
        layer_01_geometry = sb.GeoTIFFGeometryImported(-999)
        layer_01_display = sb.GeoTIFFDisplay(['#340042','#1E7F7A','#FCE51E'],0.8)
        super().add_layer(sb.LayerGeoTIFFDef('layer_01','Carga hidráulica aguas subterráneas','',layer_01_geometry,1,False,False,layer_01_display))
        super().add_layer_data(sb.LayerDataTimeSeriesDef('layer_01',sb.ChartDisplay(temporal_units_label,'Metros sobre nivel del mar',output_decimal_places)))

        # Layer 02 DTWT (GeoTIFF)
        layer_02_geometry = sb.GeoTIFFGeometryImported(-999)
        layer_02_display = sb.GeoTIFFDisplay(['#6D010E','#FFFFFF','#000000'],0.8,[-100,100])
        super().add_layer(sb.LayerGeoTIFFDef('layer_02','Profundidad a napa freática','',layer_02_geometry,2,False,False,layer_02_display))
        super().add_layer_data(sb.LayerDataTimeSeriesDef('layer_02',sb.ChartDisplay(temporal_units_label,'Metros bajo terreno',output_decimal_places)))

        # Layer 03 Recovery (GeoTIFF)
        layer_03_geometry = sb.GeoTIFFGeometryImported(-999)
        layer_03_display = sb.GeoTIFFDisplay(['#6D010E','#FFFFFF'],0.8)
        super().add_layer(sb.LayerGeoTIFFDef('layer_03','Recuperación','',layer_03_geometry,3,False,False,layer_03_display))
        super().add_layer_data(sb.LayerDataTimeSeriesDef('layer_03',sb.ChartDisplay(temporal_units_label,'Metros sobre terreno',output_decimal_places)))

        # Layer 04 Saturated Thickness (GeoTIFF)
        layer_04_geometry = sb.GeoTIFFGeometryImported(-999)
        layer_04_display = sb.GeoTIFFDisplay(['#340042','#1E7F7A','#FCE51E'],0.8)
        super().add_layer(sb.LayerGeoTIFFDef('layer_04','Espesor saturado','',layer_04_geometry,4,False,False,layer_04_display))
        super().add_layer_data(sb.LayerDataTimeSeriesDef('layer_04',sb.ChartDisplay(temporal_units_label,'Metros',output_decimal_places)))

        # Layer 05 Pumping Costs (GeoTIFF)
        layer_05_geometry = sb.GeoTIFFGeometryImported(-999)
        layer_05_display = sb.GeoTIFFDisplay(['#00FF00','#FF0000'],0.8)
        super().add_layer(sb.LayerGeoTIFFDef('layer_05','Costos de bombeo','',layer_05_geometry,5,False,False,layer_05_display))
        super().add_layer_data(sb.LayerDataTimeSeriesDef('layer_05',sb.ChartDisplay(temporal_units_label,'CLP/1000m3',output_decimal_places)))

        # Layer 06 Improvement over baseline (GeoTIFF)
        layer_06_geometry = sb.GeoTIFFGeometryImported(-999)
        layer_06_display = sb.GeoTIFFDisplay(['#FF0000','#FFFFFF', '#00FF00'],0.8, [-50.0, 50.0])
        super().add_layer(sb.LayerGeoTIFFDef('layer_06','Simulado menos base','',layer_06_geometry,6,False,False,layer_06_display))

        # Layer 801 Sectores Acuíferos (ShapeFile)
        layer_801_custom_polygon_displays = {
            'Sector 1': sb.PolygonDisplay(1,1,0.5,'#00332F','#00332F'),
            'Sector 2': sb.PolygonDisplay(1,1,0.5,'#E10976','#E10976'),
            'Sector 3': sb.PolygonDisplay(1,1,0.5,'#339966','#339966'),
            'Sector 4': sb.PolygonDisplay(1,1,0.5,'#088EDF','#088EDF'),
            'Sector 5': sb.PolygonDisplay(1,1,0.5,'#8140FF','#8140FF'),
            'Sector 6': sb.PolygonDisplay(1,1,0.5,'#F27C1A','#F27C1A')
        }
        layer_801_display = sb.ShapefileDisplay(featureDisplayKey='Sector_DGA',customFeatureDisplays=layer_801_custom_polygon_displays)
        super().add_layer(sb.LayerShapefileDef('layer_801','Sectores Acuíferos','',11,True,False,layer_801_display))

        # Layer 802 Distritos Riego (ShapeFile)
        layer_802_default_polygon_display = sb.PolygonDisplay(1,1,0.7,'#000000','#000000')
        layer_802_custom_polygon_displays = {
            'P': sb.PolygonDisplay(1,1,0.7,'#7AE27A','#7AE27A'),
            'M': sb.PolygonDisplay(1,1,0.7,'#EAF424','#EAF424'),
            'C': sb.PolygonDisplay(1,1,0.7,'#57A0E5','#57A0E5')
        }
        layer_802_display = sb.ShapefileDisplay(polygonDisplay=layer_802_default_polygon_display,featureDisplayKey='Fuen_Riego',customFeatureDisplays=layer_802_custom_polygon_displays)
        super().add_layer(sb.LayerShapefileDef('layer_802','Distritos Riego','',12,True,False,layer_802_display))

        # Layer 803 Canales (ShapeFile)
        layer_803_custom_line_display = sb.LineDisplay(1, 1, '#64CCEF')
        layer_803_display = sb.ShapefileDisplay(None, layer_803_custom_line_display, None, '',{})
        super().add_layer(sb.LayerShapefileDef('layer_803','Canales','',13,True,False,layer_803_display))

        # Layer 804 Bocatomas Principales (ShapeFile)
        layer_804_custom_point_display = sb.PointDisplay(4, 1, 1, 1, '#5E5E5E', '#EFA658')
        layer_804_display = sb.ShapefileDisplay(layer_804_custom_point_display, None, None, '',{})
        super().add_layer(sb.LayerShapefileDef('layer_804','Bocatomas Principales','',14,True,False,layer_804_display))

        # Layer 805 Fluviometricas DGA (ShapeFile)
        layer_805_custom_point_display = sb.PointDisplay(4, 1, 1, 1, '#5E5E5E', '#8746C4')
        layer_805_display = sb.ShapefileDisplay(layer_805_custom_point_display, None, None, '',{})
        super().add_layer(sb.LayerShapefileDef('layer_805','Fluviométricas DGA','',15,True,False,layer_805_display))

        # Layer 806 Pozos DGA (ShapeFile)
        layer_806_custom_point_display = sb.PointDisplay(4, 1, 1, 1, '#5E5E5E', '#000000')
        layer_806_display = sb.ShapefileDisplay(layer_806_custom_point_display, None, None, '',{})
        super().add_layer(sb.LayerShapefileDef('layer_806','Pozos DGA','',16,True,False,layer_806_display))

        # Layer 807 Pozos Agricolas (ShapeFile)
        layer_807_custom_point_display = sb.PointDisplay(4, 1, 1, 1, '#5E5E5E', '#69DB67')
        layer_807_display = sb.ShapefileDisplay(layer_807_custom_point_display, None, None, '',{})
        super().add_layer(sb.LayerShapefileDef('layer_807','Pozos Agrícolas','',17,True,False,layer_807_display))

        # Layer 808 Pozos Mineria (ShapeFile)
        layer_808_custom_point_display = sb.PointDisplay(4, 1, 1, 1, '#5E5E5E', '#F2F453')
        layer_808_display = sb.ShapefileDisplay(layer_808_custom_point_display, None, None, '',{})
        super().add_layer(sb.LayerShapefileDef('layer_808','Pozos Minería','',18,True,False,layer_808_display))

        # Layer 809 Pozos Sanitaria (ShapeFile)
        layer_809_custom_point_display = sb.PointDisplay(4, 1, 1, 1, '#5E5E5E', '#5F86E2')
        layer_809_display = sb.ShapefileDisplay(layer_809_custom_point_display, None, None, '',{})
        super().add_layer(sb.LayerShapefileDef('layer_809','Pozos Sanitaria','',19,True,False,layer_809_display))

        # Layer 10 Flow at DGA monitoring stations (Points)
        layer_10_display = sb.PointDisplay(radius=4, weight=1, opacity=1, fillOpacity=1, color='#5E5E5E', fillColor='#0000FF')
        super().add_layer(sb.LayerPointsDef('layer_10','Flujos medidos','',20,True,False,layer_10_display))
        super().add_layer_data(sb.LayerDataTimeSeriesDef('layer_10',sb.ChartDisplay(temporal_units_label,flow_units_label,output_decimal_places), 'value'))

        # Layer 20 Ecological Flows (Points)
        layer_20_display = sb.PointDisplay(radius=4, weight=1, opacity=1, fillOpacity=1, color='#5E5E5E', fillColor='#00FF00')
        super().add_layer(sb.LayerPointsDef('layer_20','Flujos ecológicos','',21,True,False,layer_20_display))
        super().add_layer_data(sb.LayerDataImageDef('layer_20','value'))

        # Layer 30 Water Balance at Lautaro Dam (Points)
        layer_30_display = sb.PointDisplay(radius=4, weight=1, opacity=1, fillOpacity=1, color='#5E5E5E', fillColor='#FF0000')
        super().add_layer(sb.LayerPointsDef('layer_30','Balance hídrico en Embalse Lautaro','',22,True,False,layer_30_display))
        super().add_layer_data(sb.LayerDataTimeSeriesDef('layer_30',sb.ChartDisplay(temporal_units_label,None,output_decimal_places), 'value'))

    def metadata(self):
        metadata = {}
        metadata['id'] = 'scenario_c'
        metadata['name'] = 'SimCopiapo V1.0'
        with open(os.path.join(self.static_dir, 'html/about.html'), 'r') as f:
            metadata['description'] = f.read()
        metadata['view'] = { 'lat': -27.60, 'lon': -70.43, 'zoom': 10 }
        metadata['timesteps'] = { 'steps': [25], 'unit_singular': 'Año', 'unit_plural': 'Años' }
        metadata['exports_results'] = True
        return metadata

    def load_constant_layers(self):

        self.set_layer('layer_801', sb.LayerShapefileVal(os.path.join(self.layer_dir,'sectores_acuiferos.shp'),4326,['Sector_DGA'],['Sector_DGA']))
        props = ['Predio','Cultivo','Tipo_Riego','Fuen_Riego','SectorRieg','Distrito_r','demanda_ri','perdidas_r','perdidas_c','demanda_br']
        self.set_layer('layer_802', sb.LayerShapefileVal(os.path.join(self.layer_dir,'distritos_riego.shp'),4326,['Predio'], props))
        self.set_layer('layer_803', sb.LayerShapefileVal(os.path.join(self.layer_dir,'canales.shp'),4326,['Nombre'],['Nombre']))
        self.set_layer('layer_804', sb.LayerShapefileVal(os.path.join(self.layer_dir,'bocatomas_principales.shp'),4326,['Canal'],['Canal']))
        self.set_layer('layer_805', sb.LayerShapefileVal(os.path.join(self.layer_dir,'fluviometricas_DGA.shp'),4326,['NOMBRE'],['NOMBRE']))
        self.set_layer('layer_806', sb.LayerShapefileVal(os.path.join(self.layer_dir,'pozos_DGA.shp'),4326,['NOMBRE'],['NOMBRE']))
        self.set_layer('layer_807', sb.LayerShapefileVal(os.path.join(self.layer_dir,'pozos_agricolas.shp'),24879,['id'],['id']))
        self.set_layer('layer_808', sb.LayerShapefileVal(os.path.join(self.layer_dir,'pozos_mineria.shp'),4326,['WellName'],['WellName']))
        self.set_layer('layer_809', sb.LayerShapefileVal(os.path.join(self.layer_dir,'pozos_sanitaria.shp'),4326,['Nombre'],['Nombre']))

        points_10 = [[-28.00227,-69.97823],[-27.97788,-69.99981],[-28.00227,-69.97823],[-27.51865,-70.26610],[-27.51894,-70.26604],[-27.51865,-70.26610],[-27.32316,-70.84007]]
        names_10 = ['Copiapo en pastillo', 'Copiapo en Lautaro', 'Copiapo en La Puerta', 'Mal Paso en Canal', 'Copiapo en Mal Paso', 'Copiapo en Ciudad', 'Copiapo en Angostura']
        self.set_layer('layer_10', sb.LayerPointsVal(points_10, names_10))

        points_20 = [[-27.32316,-70.84007]]
        names_20 = ['Flujos ecológicos del sector 6']
        self.set_layer('layer_20', sb.LayerPointsVal(points_20, names_20))

        points_30 = [[-27.98060,-69.99798],[-27.98460,-69.99798],[-27.98860,-69.99798],[-27.97752,-69.99844]]
        names_30 = ['volumen embalsado (m)', 'capacidad (%)', 'perdidas infiltracion (L/s)', 'descarga vertedero (L/s)']
        self.set_layer('layer_30', sb.LayerPointsVal(points_30, names_30))

    def initialise(self):
        pass

    def reset(self):
        self.set_run_status(0.0, 'Inicializando...')
        self.report = None

    def run_time_steps(self, dt, steps):
        self.run_model()

    def run_model(self):

        # ***** Get input variables *****

        Q_Pastillo_ts_name = self.get_input('input_0')

        ST_swap_CAN_ACH = self.get_input('input_1a01')
        ST_swap_CAN_ACH_option = self.get_input('input_1a02')
        ST_swap_CAN_ACH_PCwells = ST_swap_CAN_ACH and ST_swap_CAN_ACH_option == 'PC'
        ST_swap_CAN_ACH_PUwells = ST_swap_CAN_ACH and ST_swap_CAN_ACH_option == 'PU'

        ST_swap_CAS_RAM = self.get_input('input_1b01')

        ST_swap_89 = self.get_input('input_1c01')
        ST_swap_89_option = self.get_input('input_1c02')
        ST_swap_89_a_17 = ST_swap_89 and ST_swap_89_option == 'A D17'
        ST_swap_89_a_ACH = ST_swap_89 and ST_swap_89_option == 'A Aguas Chañar'
        ST_swap_89_a_S56_rio = ST_swap_89 and ST_swap_89_option == 'A S5 y S6 por rio'
        ST_swap_89_a_S5_tubo_rio = ST_swap_89 and ST_swap_89_option == 'A S5 por tuberia (excedente a rio)'
        ST_swap_89_a_S5_tubo_mar = ST_swap_89 and ST_swap_89_option == 'A S5 por tuberia (excedente a recarga artificial)'

        ST_lautaro2 = self.get_input('input_2a01')

        ST_entubamiento_canales = self.get_input('input_2b01')

        y_PC01 = self.get_input('input_2c01')
        y_PC02 = self.get_input('input_2c02')
        y_PC03 = self.get_input('input_2c03')
        y_PC04 = self.get_input('input_2c04')
        y_PC05 = self.get_input('input_2c05')
        y_PC06 = self.get_input('input_2c06')
        y_PC07 = self.get_input('input_2c07')
        y_PU01 = self.get_input('input_2c08')
        y_PU02 = self.get_input('input_2c09')
        y_PU03 = self.get_input('input_2c10')
        y_CE01 = self.get_input('input_2c11')

        ST_recarga_artificial_S3Nantoco = self.get_input('input_3a01')
        ST_recarga_artificial_S4AntesKaukari = self.get_input('input_3a02')
        ST_recarga_artificial_S4DespuesKaukari = self.get_input('input_3a03')
        ST_recarga_artificial_S5PiedraColgada = self.get_input('input_3a04')
        nominal_MAR_rate = self.get_input('input_3a05')
        # Available areas [m2]
        MAR_area_S3Nantoco = 35000.0
        MAR_area_S4AntesKaukari = 35000.0
        MAR_area_S4DespuesKaukari = 35000.0
        MAR_area_S5PiedraColgada = 35000.0
        # Nominal MAR capacities [L/s]
        if nominal_MAR_rate > 0:
            MAR_capacity_S3Nantoco = nominal_MAR_rate * MAR_area_S3Nantoco * 1000.0 / 86400.0
            MAR_capacity_S4AntesKaukari = nominal_MAR_rate * MAR_area_S4AntesKaukari * 1000.0 / 86400.0
            MAR_capacity_S4DespuesKaukari = nominal_MAR_rate * MAR_area_S4DespuesKaukari * 1000.0 / 86400.0
            MAR_capacity_S5PiedraColgada = nominal_MAR_rate * MAR_area_S5PiedraColgada * 1000.0 / 86400.0
        else:
            MAR_capacity_S3Nantoco = 0.0
            MAR_capacity_S4AntesKaukari = 0.0
            MAR_capacity_S4DespuesKaukari = 0.0
            MAR_capacity_S5PiedraColgada = 0.0

        prorrata_S2_agricola = self.get_input('input_4a01')
        prorrata_S3_agricola = self.get_input('input_4a02')
        prorrata_S4_agricola = self.get_input('input_4a03')
        prorrata_S5_agricola = self.get_input('input_4a04')
        prorrata_S6_agricola = self.get_input('input_4a05')

        prorrata_S2_minero = self.get_input('input_4a06')
        prorrata_S3_minero = self.get_input('input_4a07')
        prorrata_S4_minero = self.get_input('input_4a08')
        prorrata_S5_minero = self.get_input('input_4a09')
        prorrata_S6_minero = self.get_input('input_4a10')

        prorrata_M = self.get_input('input_4a11')
        prorrata_F = self.get_input('input_4a12')

        reduccion_consumo_AP = self.get_input('input_5a01') / 100.0

        # ***** Run surface water model *****

        self.set_run_status(5.0, 'Ejecutando Modelo de Aguas Superficiales...')

        # Load precomputed results from NetLogo (ABM)

        def compute_ABM_result(M,F):

            ABM_results = np.load(os.path.join(self.model_data_dir, 'resultados_abm.npy'))

            if M == 0.1: index_row = 1
            if M == 0.2: index_row = 2
            if M == 0.3: index_row = 3    
            if M == 0.4: index_row = 4    
            if M == 0.5: index_row = 5
            if M == 0.6: index_row = 6        
            if M == 0.7: index_row = 7        
            if M == 0.8: index_row = 8        
            if M == 0.9: index_row = 9        

            if F == 0.1: index_col = 1
            if F == 0.2: index_col = 2
            if F == 0.3: index_col = 3    
            if F == 0.4: index_col = 4    
            if F == 0.5: index_col = 5
            if F == 0.6: index_col = 6        
            if F == 0.7: index_col = 7        
            if F == 0.8: index_col = 8        
            if F == 0.9: index_col = 9

            return ABM_results[index_row][index_col] / 100

        if (prorrata_S2_agricola != 1 or
            prorrata_S3_agricola != 1 or
            prorrata_S4_agricola != 1 or
            prorrata_S5_agricola != 1 or
            prorrata_S6_agricola != 1):
            compliance = compute_ABM_result(prorrata_M, prorrata_F)
        else:
            compliance = 0

        # Run Lautaro Dam operations submodel
        sw_output = dam_ops.run_dam_operation_model(Q_Pastillo_ts_name, ST_swap_CAS_RAM, ST_lautaro2)
        ModeloEmbalseLautaro_df_6h = sw_output[0]
        LaPuerta_GWSW_df_6h = sw_output[1]

        # Resample submodel from 6-hour to monthly timesteps
        ModeloEmbalseLautaro_df_M = ModeloEmbalseLautaro_df_6h.resample('M').mean()
        LaPuerta_GWSW_df_M = LaPuerta_GWSW_df_6h.resample('M').mean()

        # Apply rolling filter to Q_Afloramiento and then recalculate Q_LaPuerta_sim
        a = LaPuerta_GWSW_df_M.Q_Afloramiento
        b = LaPuerta_GWSW_df_M.Q_Afloramiento.rolling(20, center=True).mean()
        b = b.fillna(b.mean())
        LaPuerta_GWSW_df_M['Q_Afloramiento'] = b
        LaPuerta_GWSW_df_M['Q_LaPuerta_sim'] = LaPuerta_GWSW_df_M['Q_Afloramiento'] + ModeloEmbalseLautaro_df_M['Q_Lautaro_sim'] + ModeloEmbalseLautaro_df_M['Q_Vertedero_sim']

        # Load base irrigation demands for JVRC irrigation districts (Sectors 2, 3, 4)
        irrigationdemands_S234_df = pd.read_pickle(os.path.join(self.model_data_dir, 'demandas_riego_S234.pkl'))

        # Load base irrigation demands for GW irrigation districts (Sectors 5, 6)
        irrigationdemands_S56_df = pd.read_pickle(os.path.join(self.model_data_dir, 'demandas_riego_S56.pkl'))

        # Load seasonal demand curves
        demandcurves_df = pd.read_pickle(os.path.join(self.model_data_dir, 'curvas_demandas_riego.pkl'))

        # Run surface water delivery submodel

        from swmodel import run_swmodel
        SWMODEL_out_df = run_swmodel(
            prorrata_S2_agricola, prorrata_S3_agricola, prorrata_S4_agricola, prorrata_S5_agricola, prorrata_S6_agricola, compliance,
            ST_swap_89_a_17, ST_swap_89_a_ACH, ST_swap_89_a_S56_rio, ST_swap_89_a_S5_tubo_rio, ST_swap_89_a_S5_tubo_mar,
            ST_lautaro2, ST_entubamiento_canales,
            ST_recarga_artificial_S3Nantoco, MAR_capacity_S3Nantoco, ST_recarga_artificial_S4AntesKaukari, MAR_capacity_S4AntesKaukari,
            ST_recarga_artificial_S4DespuesKaukari, MAR_capacity_S4DespuesKaukari, ST_recarga_artificial_S5PiedraColgada, MAR_capacity_S5PiedraColgada,
            ModeloEmbalseLautaro_df_M, LaPuerta_GWSW_df_M,
            irrigationdemands_S234_df, irrigationdemands_S56_df, demandcurves_df)

        # Change dates and trim SW model output to match simulation period
        old_sw_date_index = pd.date_range(start=SWMODEL_out_df.index[0], periods=len(SWMODEL_out_df.index), freq='M', name='date')
        new_sw_date_index = pd.date_range(start=self._start, periods=len(SWMODEL_out_df.index), freq='M', name='date')
        SWMODEL_out_df.index = new_sw_date_index

        #SWMODEL_out_df.to_pickle(os.path.join(self.model_output_dir, 'SW_model_outputs.pkl'))

        # @@@@@

        dates = SWMODEL_out_df.index.to_pydatetime()

        self.set_layer_data('layer_10', 'Copiapo en pastillo', sb.TimseriesBaseVal(dates, SWMODEL_out_df['Q_Pastillo'].values))
        self.set_layer_data('layer_10', 'Copiapo en Lautaro', sb.TimseriesBaseVal(dates, SWMODEL_out_df['Q_LautaroControlada'].values))
        self.set_layer_data('layer_10', 'Copiapo en La Puerta', sb.TimseriesBaseVal(dates, SWMODEL_out_df['Q_LaPuerta'].values))
        self.set_layer_data('layer_10', 'Mal Paso en Canal', sb.TimseriesBaseVal(dates, SWMODEL_out_df['Q_canal_in_D7'].values))
        self.set_layer_data('layer_10', 'Copiapo en Mal Paso', sb.TimseriesBaseVal(dates, SWMODEL_out_df['Q_rio_Nantoco'].values))
        self.set_layer_data('layer_10', 'Copiapo en Ciudad', sb.TimseriesBaseVal(dates, SWMODEL_out_df['Q_rio_CCopiapo'].values))
        self.set_layer_data('layer_10', 'Copiapo en Angostura', sb.TimseriesBaseVal(dates, SWMODEL_out_df['Q_rio_Angostura'].values))

        self.set_layer_data('layer_30', 'volumen embalsado (m)', sb.TimseriesBaseVal(dates, SWMODEL_out_df['Vi'].values))
        self.set_layer_data('layer_30', 'capacidad (%)', sb.TimseriesBaseVal(dates, SWMODEL_out_df['P_full'].values))
        self.set_layer_data('layer_30', 'perdidas infiltracion (L/s)', sb.TimseriesBaseVal(dates, SWMODEL_out_df['Q_InfiltracionLautaro'].values))
        self.set_layer_data('layer_30', 'descarga vertedero (L/s)', sb.TimseriesBaseVal(dates, SWMODEL_out_df['Q_Vertedero'].values))

        self.set_output('output_10', sb.TimseriesBaseVal(dates, SWMODEL_out_df['RCH_riegoycanales_S2'].values))
        self.set_output('output_11', sb.TimseriesBaseVal(dates, SWMODEL_out_df['RCH_riegoycanales_S3'].values))
        self.set_output('output_12', sb.TimseriesBaseVal(dates, SWMODEL_out_df['RCH_riegoycanales_S4'].values))
        self.set_output('output_13', sb.TimseriesBaseVal(dates, SWMODEL_out_df['Q_perdidariego_P_S5'].values))
        self.set_output('output_14', sb.TimseriesBaseVal(dates, SWMODEL_out_df['Q_perdidariego_P_S6'].values))

        self.set_output('output_20', sb.TimseriesBaseVal(dates, SWMODEL_out_df['RCH_rio_S2'].values))
        self.set_output('output_21', sb.TimseriesBaseVal(dates, SWMODEL_out_df['RCH_rio_S3'].values))
        self.set_output('output_22', sb.TimseriesBaseVal(dates, SWMODEL_out_df['RCH_rio_S4'].values))
        self.set_output('output_23', sb.TimseriesBaseVal(dates, SWMODEL_out_df['RCH_rio_S5'].values))
        self.set_output('output_24', sb.TimseriesBaseVal(dates, SWMODEL_out_df['RCH_rio_S6'].values))

        self.set_output('output_30', sb.TimseriesBaseVal(dates, SWMODEL_out_df['Satisfaccion_dda_C_D1'].values))
        self.set_output('output_31', sb.TimseriesBaseVal(dates, SWMODEL_out_df['Satisfaccion_dda_C_D2'].values))
        self.set_output('output_32', sb.TimseriesBaseVal(dates, SWMODEL_out_df['Satisfaccion_dda_C_D3'].values))
        self.set_output('output_33', sb.TimseriesBaseVal(dates, SWMODEL_out_df['Satisfaccion_dda_C_D4'].values))
        self.set_output('output_34', sb.TimseriesBaseVal(dates, SWMODEL_out_df['Satisfaccion_dda_C_D5'].values))
        self.set_output('output_35', sb.TimseriesBaseVal(dates, SWMODEL_out_df['Satisfaccion_dda_C_D6'].values))
        self.set_output('output_36', sb.TimseriesBaseVal(dates, SWMODEL_out_df['Satisfaccion_dda_C_D7'].values))
        self.set_output('output_37', sb.TimseriesBaseVal(dates, SWMODEL_out_df['Satisfaccion_dda_C_D89'].values))

        self.set_output('output_40', sb.TimseriesBaseVal(dates, SWMODEL_out_df['Satisfaccion_dda_M_D1'].values))
        self.set_output('output_41', sb.TimseriesBaseVal(dates, SWMODEL_out_df['Satisfaccion_dda_M_D2'].values))
        self.set_output('output_42', sb.TimseriesBaseVal(dates, SWMODEL_out_df['Satisfaccion_dda_M_D3'].values))
        self.set_output('output_43', sb.TimseriesBaseVal(dates, SWMODEL_out_df['Satisfaccion_dda_M_D4'].values))
        self.set_output('output_44', sb.TimseriesBaseVal(dates, SWMODEL_out_df['Satisfaccion_dda_M_D5'].values))
        self.set_output('output_45', sb.TimseriesBaseVal(dates, SWMODEL_out_df['Satisfaccion_dda_M_D6'].values))
        self.set_output('output_46', sb.TimseriesBaseVal(dates, SWMODEL_out_df['Satisfaccion_dda_M_D7'].values))
        self.set_output('output_47', sb.TimseriesBaseVal(dates, SWMODEL_out_df['Satisfaccion_dda_M_D89'].values))

        # @@@@@

        # ***** Load and create PUMPING datasets *****

        # Irrigation

        #SWMODEL_out_df = pd.read_pickle(os.path.join(self.model_output_dir, 'SW_model_outputs.pkl'))

        # Import DICTUC well cadastre
        AG_wells_gpd = pd.read_pickle(os.path.join(self.model_data_dir, 'agricultural_well_locations.pkl'))

        # Compute irrigation demands for each Aquifer Sector

        SWMODEL_gwdemands_df = SWMODEL_out_df.filter(items=[
            'Q_dda_P_new_D1', 'Q_dda_P_new_D2', 'Q_dda_P_new_D3', 'Q_dda_P_new_D4',
            'Q_dda_P_new_D5', 'Q_dda_P_new_D6', 'Q_dda_P_new_D7', 'Q_dda_P_new_D89',
            'Q_dda_P_S5', 'Q_dda_P_S6'
        ])

        SWMODEL_gwdemands_df['Q_dda_P_S2'] = SWMODEL_gwdemands_df['Q_dda_P_new_D1'] + SWMODEL_gwdemands_df['Q_dda_P_new_D2'] + SWMODEL_gwdemands_df['Q_dda_P_new_D3']
        SWMODEL_gwdemands_df['Q_dda_P_S3'] = SWMODEL_gwdemands_df['Q_dda_P_new_D4'] + SWMODEL_gwdemands_df['Q_dda_P_new_D5'] + SWMODEL_gwdemands_df['Q_dda_P_new_D6']
        SWMODEL_gwdemands_df['Q_dda_P_S4'] = SWMODEL_gwdemands_df['Q_dda_P_new_D7'] + SWMODEL_gwdemands_df['Q_dda_P_new_D89']

        AG_gwdemands_df = SWMODEL_gwdemands_df[SWMODEL_gwdemands_df.columns[-5:]]

        AG_gwdemands_df = AG_gwdemands_df.head(300)

        cols = ['Q_dda_P_S2', 'Q_dda_P_S3', 'Q_dda_P_S4', 'Q_dda_P_S5', 'Q_dda_P_S6']

        AG_gwdemands_df = AG_gwdemands_df[cols]

        AG_gwdemands_df.columns = ['Sector 2','Sector 3','Sector 4','Sector 5','Sector 6']

        AG_gwdemands_df_1M = AG_gwdemands_df

        AG_gwdemands_df_3M = AG_gwdemands_df_1M.resample('Q').mean()

        AG_gwdemands_df_3M['Sector 5'] = AG_gwdemands_df_3M['Sector 5'] * (prorrata_S5_agricola + (1 - prorrata_S5_agricola) * (1 - compliance))
        AG_gwdemands_df_3M['Sector 5'] = AG_gwdemands_df_3M['Sector 6'] * (prorrata_S6_agricola + (1 - prorrata_S6_agricola) * (1 - compliance))

        # Define functions that compute pumping schedules for each well, and then for the whole basin

        # Assume GW demand is equally distributed across al operational irrigation wells within each Aquifer Sector

        def compute_irrigation_pumping_schedule_SECTORS(sector, AG_wells_gpd, AG_gwdemands_df):
            
            demands = AG_gwdemands_df[sector]
            pozos = AG_wells_gpd[AG_wells_gpd['Sector_DGA'] == sector][['geometry']]
            numero_pozos = pozos.shape[0]
            
            pozos['Q_schedule'] = pozos.apply(lambda r: demands.values / numero_pozos, axis=1)
            pozos['Sector'] = sector
            
            return pozos

        def compute_irrigation_pumping_schedules_BASIN():
            
            pumping_schedule_S2 = compute_irrigation_pumping_schedule_SECTORS('Sector 2', AG_wells_gpd, AG_gwdemands_df_3M)
            pumping_schedule_S3 = compute_irrigation_pumping_schedule_SECTORS('Sector 3', AG_wells_gpd, AG_gwdemands_df_3M)
            pumping_schedule_S4 = compute_irrigation_pumping_schedule_SECTORS('Sector 4', AG_wells_gpd, AG_gwdemands_df_3M)
            pumping_schedule_S5 = compute_irrigation_pumping_schedule_SECTORS('Sector 5', AG_wells_gpd, AG_gwdemands_df_3M)
            pumping_schedule_S6 = compute_irrigation_pumping_schedule_SECTORS('Sector 6', AG_wells_gpd, AG_gwdemands_df_3M)  
            
            pumping_schedule = pd.concat([pumping_schedule_S2, pumping_schedule_S3 , pumping_schedule_S4 , pumping_schedule_S5, pumping_schedule_S6])
            
            return pumping_schedule

        AG_pumpingschedule = compute_irrigation_pumping_schedules_BASIN()

        #AG_pumpingschedule.to_pickle(os.path.join(self.model_output_dir, 'irrigation_wells_schedule_SCENARIO.pkl'))

        # Mining

        MINE_wells_gdf = pd.read_pickle(os.path.join(self.model_data_dir, 'mining_wells_schedule_BASE.pkl'))

        if ST_swap_CAS_RAM == True:
            MINE_wells_gdf.at[0, 'Rate'] = 0
            MINE_wells_gdf.at[1, 'Rate'] = 0
            MINE_wells_gdf.at[2, 'Rate'] = 0
            MINE_wells_gdf.at[3, 'Rate'] = 0
            MINE_wells_gdf.at[4, 'Rate'] = -80

        MINE_wells_gdf.loc[MINE_wells_gdf['Sector_DGA'] == 'Sector 2', 'Rate'] = MINE_wells_gdf['Rate'] * prorrata_S2_minero
        MINE_wells_gdf.loc[MINE_wells_gdf['Sector_DGA'] == 'Sector 3', 'Rate'] = MINE_wells_gdf['Rate'] * prorrata_S3_minero
        MINE_wells_gdf.loc[MINE_wells_gdf['Sector_DGA'] == 'Sector 4', 'Rate'] = MINE_wells_gdf['Rate'] * prorrata_S4_minero
        MINE_wells_gdf.loc[MINE_wells_gdf['Sector_DGA'] == 'Sector 5', 'Rate'] = MINE_wells_gdf['Rate'] * prorrata_S5_minero
        MINE_wells_gdf.loc[MINE_wells_gdf['Sector_DGA'] == 'Sector 6', 'Rate'] = MINE_wells_gdf['Rate'] * prorrata_S6_minero

        #MINE_wells_gdf.to_pickle(os.path.join(self.model_output_dir, 'mining_wells_schedule_SCENARIO.pkl'))

        # Drinking Water

        # Import well locations and template for drinking water production scheme
        DW_wells_gdf = pd.read_pickle(os.path.join(self.model_data_dir, 'DW_wells_locations.pkl'))
        DW_MODEL_df = pd.read_excel(os.path.join(self.model_data_dir, 'DW model.xlsx'), sheet_name='Modelo_Produccion_AP', skiprows=2, usecols="A:P")
        DW_MODEL_df = DW_MODEL_df.head(26)

        # Create DataFrame for drinking water production and desalination scheme

        def compute_desalination_stage(Q_desaladora):
            
            if Q_desaladora == 0: etapa = 0       
            
            if (Q_desaladora > 0): 
                
                if Q_desaladora < 450: etapa = 1
                    
                else:
                    
                    if Q_desaladora < 900: etapa = 2
                
                    else: etapa = 3
            
            return etapa

        def compute_DW_scheme(y_PC01, y_PC02, y_PC03, y_PC04, y_PC05, y_PC06, y_PC07, y_PU01, y_PU02, y_PU03, y_CE01,
                              ST_swap_CAN_ACH, ST_swap_CAN_ACH_PCwells, ST_swap_CAN_ACH_PUwells, ST_swap_89_a_ACH):

            # Design rates for drinking water wells
            Q_PC01 = 70
            Q_PC02 = 71
            Q_PC03 = 67
            Q_PC04 = 72
            Q_PC05 = 85
            Q_PC06 = 74
            Q_PC07 = 77
            Q_PU01 = 68
            Q_PU02 = 68
            Q_PU03 = 68
            Q_CE01 = 50

            if ST_swap_CAN_ACH == True:
                if ST_swap_CAN_ACH_PCwells == True:
                    Q_PC01 = 0
                    Q_PC02 = 0
                    Q_PC03 = 33
                if ST_swap_CAN_ACH_PUwells == True:
                    Q_PU01 = 0
                    Q_PU02 = 0
                    Q_PU03 = 39

            DW_MODEL_df['PC01'] = np.where(DW_MODEL_df['YEAR'] >= y_PC01, 0, Q_PC01)
            DW_MODEL_df['PC02'] = np.where(DW_MODEL_df['YEAR'] >= y_PC02, 0, Q_PC02)
            DW_MODEL_df['PC03'] = np.where(DW_MODEL_df['YEAR'] >= y_PC03, 0, Q_PC03)
            DW_MODEL_df['PC04'] = np.where(DW_MODEL_df['YEAR'] >= y_PC04, 0, Q_PC04)
            DW_MODEL_df['PC05'] = np.where(DW_MODEL_df['YEAR'] >= y_PC05, 0, Q_PC05)
            DW_MODEL_df['PC06'] = np.where(DW_MODEL_df['YEAR'] >= y_PC06, 0, Q_PC06)
            DW_MODEL_df['PC07'] = np.where(DW_MODEL_df['YEAR'] >= y_PC07, 0, Q_PC07)
            DW_MODEL_df['PU01'] = np.where(DW_MODEL_df['YEAR'] >= y_PU01, 0, Q_PU01)
            DW_MODEL_df['PU02'] = np.where(DW_MODEL_df['YEAR'] >= y_PU02, 0, Q_PU02)
            DW_MODEL_df['PU03'] = np.where(DW_MODEL_df['YEAR'] >= y_PU03, 0, Q_PU03)
            DW_MODEL_df['CE01'] = np.where(DW_MODEL_df['YEAR'] >= y_CE01, 0, Q_CE01)

            if ST_swap_89_a_ACH == True:
                DW_MODEL_df['DDA EXPLOTACION'] = DW_MODEL_df['DDA EXPLOTACION'] - SWMODEL_out_df.Q_canal_in_D89.resample('Y').mean().head(26).reset_index().Q_canal_in_D89

            DW_MODEL_df['DDA EXPLOTACION'] = DW_MODEL_df['DDA EXPLOTACION'] * (1 - reduccion_consumo_AP)

            DW_MODEL_df[DW_MODEL_df < 0] = 0

            DW_MODEL_df['Desaladora'] = DW_MODEL_df['DDA EXPLOTACION'] - (\
                                                                          DW_MODEL_df['PC01'] + 
                                                                          DW_MODEL_df['PC02'] + 
                                                                          DW_MODEL_df['PC03'] + 
                                                                          DW_MODEL_df['PC04'] + 
                                                                          DW_MODEL_df['PC05'] + 
                                                                          DW_MODEL_df['PC06'] +
                                                                          DW_MODEL_df['PC07'] +
                                                                          DW_MODEL_df['PU01'] + 
                                                                          DW_MODEL_df['PU02'] + 
                                                                          DW_MODEL_df['PU03'] +
                                                                          DW_MODEL_df['CE01']
                                                                         )


            DW_MODEL_df['Desaladora'] = np.where(DW_MODEL_df['Desaladora'] <= 0, 0, DW_MODEL_df['Desaladora'])

            DW_MODEL_df['Etapa Desaladora']=DW_MODEL_df['Desaladora'].apply(lambda x: compute_desalination_stage(x))

            DW_MODEL_df['Perdidas'] = DW_MODEL_df['DDA EXPLOTACION'] * 0.3

            return DW_MODEL_df

        # Run DW submodel
        DW_MODEL_df = compute_DW_scheme(y_PC01, y_PC02, y_PC03, y_PC04, y_PC05, y_PC06, y_PC07, y_PU01, y_PU02, y_PU03, y_CE01,
                                        ST_swap_CAN_ACH, ST_swap_CAN_ACH_PCwells, ST_swap_CAN_ACH_PUwells, ST_swap_89_a_ACH)

        # @@@@@

        # Desalination plant

        self.set_output('output_50', sb.OutputMultipleNumericDatasetVal().append_dataset([int(x) for x in DW_MODEL_df['Desaladora'].head(25).values]))

        # @@@@@

        # Resample to quaterly timesteps
        DW_MODEL_df['YEAR'] = pd.to_datetime(DW_MODEL_df['YEAR'], format='%Y')
        DW_MODEL_df = DW_MODEL_df.set_index('YEAR')
        DW_MODEL_df_3M = DW_MODEL_df.resample('Q').pad().reset_index().drop(columns='YEAR')
        DW_MODEL_df_3M = DW_MODEL_df_3M.astype(int)
        DW_MODEL_df_3M = DW_MODEL_df_3M.head(100)

        #DW_MODEL_df_3M.to_pickle(os.path.join(self.model_output_dir, 'drinking_wells_schedule_SCENARIO.pkl'))

        # ***** Load and create RECHARGE datasets *****

        # Import GeoDataframes with irrigation, river and city cells
        model_irrcells_gdf = pd.read_pickle(os.path.join(self.model_data_dir, 'model_irrcells.pkl'))
        model_rivcells_gdf = pd.read_pickle(os.path.join(self.model_data_dir, 'model_rivcells.pkl'))
        model_citcells_gdf = pd.read_pickle(os.path.join(self.model_data_dir, 'model_citcells.pkl'))

        # Compute irrigation (infiltration + canal) RCH fluxes per aquifer sector
        RCH_irr_df_1M = SWMODEL_out_df[['RCH_riegoycanales_S2', 'RCH_riegoycanales_S3', 'RCH_riegoycanales_S4', 'Q_perdidariego_P_S5', 'Q_perdidariego_P_S6']]
        RCH_irr_df_3M = RCH_irr_df_1M.resample('Q').mean()
        RCH_irr_df_3M = RCH_irr_df_3M.reset_index().drop(columns='date').head(100)
        RCH_irr_df_3M.columns = ['Sector 2', 'Sector 3', 'Sector 4', 'Sector 5', 'Sector 6']
        RCH_irr_df_3M['Sector 5'] = RCH_irr_df_3M['Sector 5'] * (prorrata_S5_agricola + (1 - prorrata_S5_agricola) * (1 - compliance))
        RCH_irr_df_3M['Sector 6'] = RCH_irr_df_3M['Sector 6'] * (prorrata_S6_agricola + (1 - prorrata_S6_agricola) * (1 - compliance))

        # Compute Copiapo River RCH fluxes per aquifer sector
        RCH_riv_df_1M = SWMODEL_out_df[['RCH_rio_S2', 'RCH_rio_S3', 'RCH_rio_S4', 'RCH_rio_S5', 'RCH_rio_S6']]
        RCH_riv_df_3M = RCH_riv_df_1M.resample('Q').mean()
        RCH_riv_df_3M = RCH_riv_df_3M.reset_index().drop(columns='date').head(100)
        RCH_riv_df_3M.columns = ['Sector 2','Sector 3','Sector 4','Sector 5','Sector 6']

        # Compute Lautaro Dam RCH fluxes
        model_laucells_gdf = gpd.read_file(os.path.join(self.model_geodata_dir, 'Lautaro_recharge.shp'))[['row','column','geometry']]
        RCH_lau_df_1M = SWMODEL_out_df['Q_InfiltracionLautaro']
        RCH_lau_df_3M = RCH_lau_df_1M.resample('Q').mean().reset_index().drop(columns='date').head(100)
        RCH_lau_df_3M[RCH_lau_df_3M < 0] = 0     # replace negative values with 0

        # Compute drinking water network loss fluxes
        RCH_cit_df_3M = DW_MODEL_df_3M.head(100)

        # ***** Load base ground water model *****

        self.set_run_status(15.0, 'Cargando Modelo de Aguas Subterráneas...')

        def load_base_groundwater_model():

            exe_name = os.path.join(self.model_modflow_bin_dir, 'mfnwt')
            version = 'mfnwt'

            import zipfile
            zf = zipfile.ZipFile(os.path.join(self.model_modflow_data_dir, 'load.zip'), 'r')
            zf.extractall(self.model_modflow_load_dir)
            zf.close()

            loadpath = self.model_modflow_load_dir
            runpath = self.model_modflow_run_dir

            # Read MODFLOW model, set model name and create new path to store simulation files
            ml = flopy.modflow.Modflow.load('SIMCOPIAPO.nam', model_ws=loadpath, version=version, exe_name=exe_name, load_only=['dis','bas6','drn','evt','nwt','wel','rch','oc','upw'])
            ml.name='SIMCOPIAPO'
            ml.change_model_ws(new_pth=runpath)

            # Tweak number of timesteps
            ml.dis.nstp = 1
            # 3-month stress periods
            ml.dis.perlen = 30.4 * 3

            # Use starting heads from Hidromas model (comment out if using original strt surface)
            ml.bas6.strt = np.load(os.path.join(self.model_data_dir, 'strt.npy'))

            # Remove .hds file reference from the namefile
            ml.remove_output(unit=51)
            ml.remove_external(unit=51)

            # Save model outputs every 6 months — January (summer) + June (winter)
            spd = {}
            for i in range(0,99,2):
                spd[(i,0)] = ['save head', 'save budget']

            oc = flopy.modflow.mfoc.ModflowOc(ml, stress_period_data=spd)

            # Make sure spatial reference set to epsg 24879 to create RCH and WEL packages
            xul = 309000
            yul = 6993800
            rot = 0
            model_epsg = 24879
            ml.sr = flopy.utils.reference.SpatialReference(delr=ml.dis.delr.array, delc=ml.dis.delc.array, xul=xul, yul=yul, rotation=rot, epsg=model_epsg)

            return ml

        ml = load_base_groundwater_model()

        # ***** Write RCH, WEL packages for ground water model for run *****

        self.set_run_status(30.0, 'Preparando Paquetes del Modelo de Aguas Subterráneas...')

        # *** Write RCH package ***

        def threaded_rch_spd_calc():

            pool = Pool(processes=4)

            workers = []
            for i in range(5):
                sp = np.arange(i*20,i*20+20,1)
                args = (model_irrcells_gdf, model_rivcells_gdf, model_laucells_gdf, model_citcells_gdf, RCH_irr_df_3M, RCH_riv_df_3M, RCH_lau_df_3M, RCH_cit_df_3M, sp, )
                workers.append(pool.apply_async(t_funcs.rch, args=args))

            pool.close()
            pool.join()

            rch_spd = {}

            for i in range(len(workers)):
                r = workers[i].get()
                rch_spd.update(r)

            pool = None

            return rch_spd

        rch_spd = threaded_rch_spd_calc()

        #pickle.dump(rch_spd, open(os.path.join(self.model_output_dir, 'rch_spd_SCENARIO.pkl'), "wb" ))

        rch = flopy.modflow.ModflowRch(ml, nrchop=3, rech=rch_spd)

        # *** Write WEL package ***

        # Irrigation

        def threaded_irr_well_sp_calc():

            irrigation_wells_schedule = AG_pumpingschedule
            irrigation_wells = irrigation_wells_schedule.index.values

            # Pre-extract row, col for irrigation wells
            rcs = {}
            for well in irrigation_wells:
                row = ml.sr.get_rc(irrigation_wells_schedule.geometry.x[well], irrigation_wells_schedule.geometry.y[well])[0]
                col = ml.sr.get_rc(irrigation_wells_schedule.geometry.x[well], irrigation_wells_schedule.geometry.y[well])[1]
                rcs[well] = (row,col)

            pool = Pool(processes=4)

            workers = []
            for i in range(5):
                sp = np.arange(i*20,i*20+20,1)
                args = (rcs, irrigation_wells_schedule, sp, irrigation_wells, )
                workers.append(pool.apply_async(t_funcs.irr_wel_sp, args=args))

            pool.close()
            pool.join()

            sp_dict_irr = {}

            for i in range(len(workers)):
                r = workers[i].get()
                sp_dict_irr.update(r)

            pool = None

            return sp_dict_irr

        sp_dict_IRR = threaded_irr_well_sp_calc()

        #pickle.dump(sp_dict_IRR, open(os.path.join(self.model_output_dir, 'sp_dict_IRR.pkl'),'wb'))

        wel_IRR = flopy.modflow.ModflowWel(ml, stress_period_data = sp_dict_IRR)

        # Mining

        def threaded_min_well_sp_calc():

            mine_wells_gdf = MINE_wells_gdf
            mine_wells = mine_wells_gdf.index.values

            # Pre-extract row, col for mine wells
            rcs = {}
            for well in mine_wells:
                row = ml.sr.get_rc(mine_wells_gdf.geometry.x[well], mine_wells_gdf.geometry.y[well])[0]
                col = ml.sr.get_rc(mine_wells_gdf.geometry.x[well], mine_wells_gdf.geometry.y[well])[1]
                rcs[well] = (row,col)

            pool = Pool(processes=4)

            workers = []
            for i in range(5):
                sp = np.arange(i*20,i*20+20,1)
                args = (rcs, mine_wells_gdf, sp, mine_wells, )
                workers.append(pool.apply_async(t_funcs.min_wel_sp, args=args))

            pool.close()
            pool.join()

            sp_dict_min = {}

            for i in range(len(workers)):
                r = workers[i].get()
                sp_dict_min.update(r)

            pool = None

            return sp_dict_min

        sp_dict_MINE = threaded_min_well_sp_calc()

        #pickle.dump(sp_dict_MINE, open(os.path.join(self.model_output_dir, 'sp_dict_MINE.pkl'),'wb'))

        wel_MINE = flopy.modflow.ModflowWel(ml, stress_period_data = sp_dict_MINE)

        # Drinking water

        def threaded_dw_well_sp_calc():

            dw_schedule = DW_MODEL_df_3M[['PC01','PC02','PC03','PC04','PC05','PC06','PC07','PU01','PU02','PU03','CE01']].reset_index().T
            dw_wells = DW_wells_gdf.index.values

            # Pre-extract row, col for dw wells
            rcs = {}
            for well in dw_wells:
                row = ml.sr.get_rc(DW_wells_gdf.geometry.x[well], DW_wells_gdf.geometry.y[well])[0]
                col = ml.sr.get_rc(DW_wells_gdf.geometry.x[well], DW_wells_gdf.geometry.y[well])[1]
                rcs[well] = (row,col)

            pool = Pool(processes=4)

            workers = []
            for i in range(5):
                sp = np.arange(i*20,i*20+20,1)
                args = (rcs, dw_schedule, sp, dw_wells, )
                workers.append(pool.apply_async(t_funcs.dw_wel_sp, args=args))

            pool.close()
            pool.join()

            sp_dict_dw = {}

            for i in range(len(workers)):
                r = workers[i].get()
                sp_dict_dw.update(r)

            pool = None

            return sp_dict_dw

        sp_dict_DW = threaded_dw_well_sp_calc()

        #pickle.dump(sp_dict_DW, open(os.path.join(self.model_output_dir, 'sp_dict_DW_rb.pkl'),'wb'))

        wel_DW = flopy.modflow.ModflowWel(ml, stress_period_data = sp_dict_DW)

        # Combine WEL_IRR + WEL_DW + WEL_MINE into single WEL package

        wel_IRR_DW = flopy.modflow.ModflowWel(ml, stress_period_data = wel_DW.stress_period_data.append(wel_IRR.stress_period_data))

        wel = flopy.modflow.ModflowWel(ml, stress_period_data = wel_IRR_DW.stress_period_data.append(wel_MINE.stress_period_data))

        # Write ground water model input

        self.set_run_status(40.0, 'Inicializando Modelo de Aguas Subterráneas...')

        ml.oc.reset_budgetunit(budgetunit=1053, fname='SIMCOPIAPO.cbc')

        ml.exe_name = os.path.join(self.model_modflow_bin_dir, 'mfnwt')
        if os.name == 'nt': # Windows
            ml.exe_name = ml.exe_name + '.exe'

        ml.write_input()

        # Run ground water model

        self.set_run_status(60.0, 'Ejecutando Modelo de Aguas Subterráneas...')

        ml.run_model()#silent=True

        # ***** Outputs *****

        self.set_run_status(80.0, 'Generando Resultados de Datos Espaciales...')

        # GW model

        headfile = flopy.utils.HeadFile(os.path.join(self.model_modflow_run_dir, 'SIMCOPIAPO.hds'), model=ml)

        def read_heads_file(headfile):
            heads = headfile.get_alldata()
            heads[heads==-999.98999023] = np.nan
            heads[heads==1.e+30] = np.nan
            heads[heads==-1.e+30] = np.nan
            return heads

        # Heads file has 100 time slices (zero-based indexing) i.e., outputs are retrieved every 3-months — odd slices = summer ; even slices = winter
        heads_3D_array = read_heads_file(headfile)

        # Compute ground water outputs

        def compute_gw_outputs(ml, heads_3D_array, time_slices):
            heads = []
            dtwt = []
            gwrecov = []
            sat_thickness = []
            for i in range(time_slices):
                heads.append(heads_3D_array[i][0])                          # Head
                dtwt.append(ml.dis.top.array - heads[i])                    # Depth to Water Table (meters)
                gwrecov.append(np.ma.masked_where(dtwt[i] >= 0, dtwt[i]))   # Areas of GW recovery
                sat_thickness.append(heads[i] - ml.dis.botm.array[0])       # Saturated thickness
            return heads, dtwt, gwrecov, sat_thickness

        time_slices = 50
        heads, dtwt, gwrecov, sat_thickness = compute_gw_outputs(ml, heads_3D_array, time_slices)

        # @@@@@

        self._apply_geotiff_layer(ml, heads, 49, 'layer_01', 'EPSG:24879', np.nan)
        self._apply_geotiff_layer(ml, dtwt, 49, 'layer_02', 'EPSG:24879', np.nan)
        self._apply_geotiff_layer(ml, gwrecov, 49, 'layer_03', 'EPSG:24879', np.nan)
        self._apply_geotiff_layer(ml, sat_thickness, 49, 'layer_04', 'EPSG:24879', np.nan)

        # @@@@@

        # Extract summer and winter heads

        heads_3D_array_summer = heads_3D_array[::2]

        heads_3D_array_winter = heads_3D_array[1::2]

        # Load groundwater heads baseline case

        heads_baseline = np.load(os.path.join(self.model_data_dir, 'baseline_heads.npy'))

        # Compute scenario improvement over baseline

        heads_scenario_improvement = heads_3D_array[49][0] - heads_baseline

        # @@@@@

        self._apply_single_geotiff_layer(ml, heads_scenario_improvement, 'layer_06', 'EPSG:24879', np.nan)

        # @@@@@

        # Store ground water model output arrays in a single shapefile

        mfg = gpd.read_file(os.path.join(self.model_geodata_dir, 'zonebudget_zones.shp'))

        def export_gridded_outputs_to_shapefile(ml, mfg, heads_3D_array):
            time_slices = np.arange(0,len(heads_3D_array),1)

            ibound = ml.bas6.ibound.array[0].astype(int)
            thickness = ml.dis.thickness.array[0]

            mfg['ibound'] = ibound.flatten()
            mfg['thickness'] = thickness.flatten()

            for time_slice in time_slices:
                mfg['heads' + '_' + str(time_slice)] = heads_3D_array[time_slice][0].flatten()
                mfg['dtwt' + '_' + str(time_slice)] = (ml.dis.top.array - heads_3D_array[time_slice][0]).flatten()
                mfg['gwrecov' + '_' + str(time_slice)] = (np.ma.masked_where(ml.dis.top.array - heads_3D_array[time_slice][0] >= 0, ml.dis.top.array - heads_3D_array[time_slice][0])).flatten()
                mfg['sat' + '_' + str(time_slice)] = (heads_3D_array[time_slice][0] - ml.dis.botm.array[0]).flatten()

            mfg_out = mfg.loc[mfg['ibound']==1]
            return mfg_out#mfg_out.to_file(os.path.join(self.model_output_dir, 'GW_model_outputs.shp'))

        GW_model_outputs = export_gridded_outputs_to_shapefile(ml, mfg, heads_3D_array_summer)

        self.set_run_status(90.0, 'Finalizando Resultados del Modelo...')

        # Compute pumping costs (GeoTIFF)

        tarriff_kwh = 80.0

        def compute_pumping_cost(ml, heads_3D_array, tarriff_kwh):
            dtwt_3D_array = ml.dis.top.array - heads_3D_array
            vec_func = np.vectorize(lambda x : x * (2.725 / 0.7 / 0.95 / 0.9 * tarriff_kwh))
            dtwt_3D_array = vec_func(dtwt_3D_array)
            return np.squeeze(dtwt_3D_array)

        pumping_costs = compute_pumping_cost(ml, heads_3D_array, tarriff_kwh)

        # @@@@@

        self._apply_geotiff_layer(ml, pumping_costs, 49, 'layer_05', 'EPSG:24879', np.nan)

        # @@@@@

        # Compute changes in aquifer storage (per sector)

        def compute_changes_in_aquifer_storage(GW_model_outputs):

            df = GW_model_outputs[(GW_model_outputs['ibound'] == 1) & (GW_model_outputs['thickness'] > 0.0)]
            df['initial_storage'] = df['sat_0'] * (200 ** 2) / 1000000.0                                            # in million cubic metres

            #5 years
            df['decrease_in_storage_5y'] = (df['sat_4'] - GW_model_outputs['sat_0']) * (200 ** 2) / 1000000.0       # in million cubic metres
            #10 years
            df['decrease_in_storage_10y'] = (df['sat_9'] - GW_model_outputs['sat_0']) * (200 ** 2) / 1000000.0      # in million cubic metres
            #15 years
            df['decrease_in_storage_15y'] = (df['sat_14'] - GW_model_outputs['sat_0']) * (200 ** 2) / 1000000.0     # in million cubic metres
            #20 years
            df['decrease_in_storage_20y'] = (df['sat_19'] - GW_model_outputs['sat_0']) * (200 ** 2) / 1000000.0     # in million cubic metres
            #25 years
            df['decrease_in_storage_25y'] = (df['sat_24'] - GW_model_outputs['sat_0']) * (200 ** 2) / 1000000.0     # in million cubic metres

            # 5 years
            df['percent_change_in_storage_5y'] = df['decrease_in_storage_5y'] / df['initial_storage'] * 100.0
            # 10 years
            df['percent_change_in_storage_10y'] = df['decrease_in_storage_10y'] / df['initial_storage'] * 100.0
            # 15 years
            df['percent_change_in_storage_15y'] = df['decrease_in_storage_15y'] / df['initial_storage'] * 100.0
            # 20 years
            df['percent_change_in_storage_20y'] = df['decrease_in_storage_20y'] / df['initial_storage'] * 100.0
            # 25 years
            df['percent_change_in_storage_25y'] = df['decrease_in_storage_25y'] / df['initial_storage'] * 100.0

            df = df.set_index('Sector_DGA')
            df.rename(index = {1.0: "Sector 1", 2.0:"Sector 2", 3.0:"Sector 3", 4.0:"Sector 4", 5.0:"Sector 5", 6.0:"Sector 6"}, inplace = True)

            # 5 years
            df['number_of_dam_volumes_5y'] = df['decrease_in_storage_5y'] / 25.0
            # 10 years
            df['number_of_dam_volumes_10y'] = df['decrease_in_storage_10y'] / 25.0
            # 15 years
            df['number_of_dam_volumes_15y'] = df['decrease_in_storage_15y'] / 25.0
            # 20 years
            df['number_of_dam_volumes_20y'] = df['decrease_in_storage_20y'] / 25.0
            # 25 years
            df['number_of_dam_volumes_25y'] = df['decrease_in_storage_25y'] / 25.0

            df = df.groupby('Sector_DGA').sum().reset_index().set_index('Sector_DGA')

            df_volumetric_change = df[['decrease_in_storage_5y','decrease_in_storage_10y','decrease_in_storage_15y','decrease_in_storage_20y','decrease_in_storage_25y']]
            df_volumetric_change = df_volumetric_change.rename(index=str, columns={"decrease_in_storage_5y": "t=5", "decrease_in_storage_10y": "t=10", "decrease_in_storage_15y": "t=15", "decrease_in_storage_20y": "t=20", "decrease_in_storage_25y": "t=25"})

            df_volumetric_change.drop('Sector 1', inplace=True)

            df_equivalent_change = df_volumetric_change / 25.0

            return (df_volumetric_change, df_equivalent_change)

        #GW_model_outputs = gpd.read_file(os.path.join(self.model_output_dir, 'GW_model_outputs.shp'))

        cas_dfs = compute_changes_in_aquifer_storage(GW_model_outputs)

        # @@@@@

        output_2_dataset = sb.OutputMultipleNumericDatasetVal()
        for k in cas_dfs[0]:
            output_2_dataset.append_dataset(list(cas_dfs[0][k].values), k)
        self.set_output('output_02', output_2_dataset)

        output_3_dataset = sb.OutputMultipleNumericDatasetVal()
        for k in cas_dfs[1]:
            output_3_dataset.append_dataset(list(cas_dfs[1][k].values), k)
        self.set_output('output_03', output_3_dataset)

        # @@@@@

        # Compute mean change in depth of water table and pumping cost per sector

        def compute_mean_change_in_depth_to_water_table(GW_model_outputs):

            df = GW_model_outputs
            df = df[(df['ibound'] == 1) & (df['thickness'] > 0.0)]
            df = df[['Sector_DGA','dtwt_0','dtwt_4','dtwt_9','dtwt_14','dtwt_19','dtwt_24']]
            df = df.groupby('Sector_DGA').mean().reset_index().set_index('Sector_DGA')
            df.rename(index = {1.0: "Sector 1", 2.0:"Sector 2", 3.0:"Sector 3", 4.0:"Sector 4", 5.0:"Sector 5", 6.0:"Sector 6"}, inplace = True) 
            df = df.rename(index=str, columns={"dtwt_0": "t=0", "dtwt_4": "t=5", "dtwt_9": "t=10", "dtwt_14": "t=15", "dtwt_19": "t=20", "dtwt_24": "t=25"})
            df.drop('Sector 1', inplace=True)
            return df

        mcdwt_df = compute_mean_change_in_depth_to_water_table(GW_model_outputs)

        # @@@@@

        output_4_dataset = sb.OutputMultipleNumericDatasetVal()
        for k in mcdwt_df:
            output_4_dataset.append_dataset(list(mcdwt_df[k].values), k)
        self.set_output('output_04', output_4_dataset)

        # @@@@@

        # Compute mean pumping costs

        def compute_mean_pumping_costs(mcdwt_df, tarriff_kwh):

            mpc_df = mcdwt_df
            for k in mpc_df:
                mpc_df[k] = mcdwt_df[k] * (2.725 / 0.7 / 0.95 / 0.9 * tarriff_kwh)            
            return mpc_df

        mpc_df = compute_mean_pumping_costs(mcdwt_df, tarriff_kwh)

        # @@@@@

        output_5_dataset = sb.OutputMultipleNumericDatasetVal()
        for k in mpc_df:
            output_5_dataset.append_dataset(list(mpc_df[k].values), k)
        self.set_output('output_05', output_5_dataset)

        # @@@@@

        # Generate ground water balance zone grid

        def threaded_zone_grid(zb_grid):

            zb_grid = zb_grid[['row','column','Sector_DGA']]
            zb_grid.fillna(value=pd.np.nan, inplace=True)
            zb_grid.replace(np.nan, 0, inplace=True)

            pool = Pool(processes=4)

            workers = []
            NUM_PER_PROC = 104
            NUM_PROC = 5
            # (104 x 5 = 520)
            for i in range(NUM_PROC):
                rows = np.arange(i*NUM_PER_PROC,i*NUM_PER_PROC+NUM_PER_PROC,1)
                args = (zb_grid, rows, i, )
                workers.append(pool.apply_async(t_funcs.zone_grid, args=args))

            pool.close()
            pool.join()

            res = np.ndarray((520,510),dtype='int')

            for i in range(len(workers)):
                r = workers[i].get()
                for j in range(NUM_PER_PROC):
                    res[r[0]*NUM_PER_PROC+j,:] = r[1][j]

            pool = None

            return res

        #zb_arr = threaded_zone_grid(mfg) # Use pre-generated one now...
        zb_arr = np.load(os.path.join(self.model_data_dir, 'zb_array.npy'))

        def generate_ground_water_zone_budget(zb_array):

            cbc_f = os.path.join(self.model_modflow_run_dir, 'SIMCOPIAPO.cbc')
            mf_list = flopy.utils.MfListBudget(os.path.join(self.model_modflow_run_dir, 'SIMCOPIAPO.list'))
            zon = zb_array.astype(int)
            aliases = {1: 'Sector 1', 2:'Sector 2', 3: 'Sector 3', 4: 'Sector 4', 5: 'Sector 5', 6: 'Sector 6'}
            zb = flopy.utils.ZoneBudget(cbc_f, zon, aliases=aliases)
            zb = zb * 1000 / 86400
            zb_df = zb.get_dataframes(start_datetime="30-9-2018", timeunit='D')
            return zb_df

        zb_df = generate_ground_water_zone_budget(zb_arr)

        # Compute ecological flow @ outlet of Sector 6

        def create_ecological_flow_at_outlet_of_sector_6_chart(zb_df, fn):

            caudal_ecologico_historico = 50
            angostura_baseflow_adjustment_factor_summer = 4
            angostura_baseflow_adjustment_factor_winter = 40

            S6_zb_df = zb_df.Sector_6.reset_index()
            S6_zb_df = S6_zb_df.loc[((S6_zb_df['name'] == 'TO_DRAINS'))]
            #mean = S6_zb_df.groupby('datetime').sum().mean().values[0]

            anf_df=S6_zb_df.groupby('datetime').sum().reset_index()
            anf_df.loc[anf_df['datetime'].dt.month == 6, ['Sector_6']] *= angostura_baseflow_adjustment_factor_summer
            anf_df.loc[anf_df['datetime'].dt.month == 12, ['Sector_6']] *= angostura_baseflow_adjustment_factor_winter
            anf_df.set_index('datetime', inplace=True)
            a=anf_df.resample('M').pad().head(300)

            aporte_subterraneo_df = a.tshift(4, freq='M')
            aporte_superficial_df = SWMODEL_out_df['Q_rio_Angostura'].tshift(12*27, freq='M').head(300)
            aporte_total = (a.Sector_6 + aporte_superficial_df)

            df_Angostura = pd.DataFrame()
            df_Angostura['aporte_superficial'] = aporte_superficial_df
            df_Angostura['aporte_subterraneo'] = aporte_subterraneo_df
            df_Angostura['total'] = aporte_total

            plt.figure(figsize=(10,5))

            plt.plot(df_Angostura.aporte_subterraneo + df_Angostura.aporte_superficial, label='superficial + subterraneo')
            plt.title('caudal humedal sector Angostura (aportes superficiales y subterraneos)', weight='bold', size=16)
            plt.ylabel('L/s', weight='bold')
            plt.legend()
            plt.ylim(0,300)
            plt.axhline(y=caudal_ecologico_historico, color='g', linestyle='-')
            plt.text('2018-6-1', caudal_ecologico_historico+10, 'caudal ecologico historico\n     (minimo) pre-1985', color='k', bbox=dict(boxstyle="round",
                               ec=(1., 0.5, 0.5),
                               fc=(1., 0.8, 0.8),))

            plt.savefig(fn)

        # Revert date index back on SWMODEL_out_df

        SWMODEL_out_df.index = old_sw_date_index

        # Set Ecological flow at outlet 6 chart

        eco_flow_outlet_6_chart_filename = os.path.join(self.model_output_dir, 'eco6out.svg')
        create_ecological_flow_at_outlet_of_sector_6_chart(zb_df, eco_flow_outlet_6_chart_filename)

        # @@@@@

        layer_20_ident = self.get_unique_image_id()
        with open(eco_flow_outlet_6_chart_filename, 'rb') as fl:
            self.set_image(layer_20_ident, fl.read(), 'image/svg+xml', True)
        self.set_layer_data('layer_20', 'Flujos ecológicos del sector 6' , sb.OutputImageVal(layer_20_ident))

        # @@@@@

        # Store required datasets for report generation

        self.report = {}

        self.report['zb_array'] = zb_arr
        self.report['zb_df'] = zb_df
        self.report['SWMODEL_out_df'] = SWMODEL_out_df
        self.report['GW_model_outputs'] = GW_model_outputs
        self.report['ST_lautaro2'] = ST_lautaro2
        self.report['DW_MODEL_df'] = DW_MODEL_df
        self.report['ml'] = ml
        self.report['headfile'] = headfile
        self.report['heads_scenario_improvement'] = heads_scenario_improvement
        self.report['compliance'] = compliance

        # Force end of simulation period
        self._now = self._end

    def generate_report(self, fp):

        # Extract required variables

        zb_array = self.report['zb_array']
        zb_df = self.report['zb_df']
        SWMODEL_out_df = self.report['SWMODEL_out_df']
        GW_model_outputs = self.report['GW_model_outputs']
        ST_lautaro2 = self.report['ST_lautaro2']
        DW_MODEL_df = self.report['DW_MODEL_df']
        ml = self.report['ml']
        headfile = self.report['headfile']
        heads_scenario_improvement = self.report['heads_scenario_improvement']
        compliance = self.report['compliance']

        # Constants

        caudal_ecologico_historico = 50
        angostura_baseflow_adjustment_factor_summer = 4
        angostura_baseflow_adjustment_factor_winter = 40

        # Groundwater recoveries Sector La Puerta (MODFLOW vs SW model)

        S2_zb_df = zb_df.Sector_2.reset_index()
        S2_zb_df = S2_zb_df.loc[(S2_zb_df['name'] == 'TO_DRAINS')]
        mean = S2_zb_df.groupby('datetime').sum().mean().values[0]

        aporte_subterraneo_df = S2_zb_df.groupby('datetime').sum().resample('M').pad().head(300)

        aporte_subterraneo_df.plot(figsize=(10,5))
        SWMODEL_out_df.Q_Afloramiento_LaPuerta.tshift(12*27, freq='M').tshift(7, freq='M').fillna(1500).head(300).plot()
        #SWMODEL_out_df.Q_Afloramiento_LaPuerta.tshift(12*27, freq='M').tshift(7, freq='M').head(300).plot()

        plt.title('Afloramiento subterraneo sector La Puerta', weight='bold', size=16)
        plt.ylabel('L/s', weight='bold')
        plt.legend(['MODFLOW','modelo aguas superficiales'])

        # Mean change in depth to water table and aquifer storage

        df1 = GW_model_outputs
        bdf1 = df1[(df1['ibound'] == 1) & (df1['thickness'] > 0)]
        df1 = df1[['Sector_DGA','dtwt_0','dtwt_4','dtwt_9','dtwt_14','dtwt_19','dtwt_24']]
        df1 = df1.groupby('Sector_DGA').mean().reset_index().set_index('Sector_DGA')
        df1.rename(index = {1.0: "Sector 1", 2.0:"Sector 2", 3.0:"Sector 3", 4.0:"Sector 4", 5.0:"Sector 5", 6.0:"Sector 6"}, inplace = True)
        df1 = df1.rename(index=str, columns={"dtwt_0": "actual", "dtwt_4": "t=5 (CP)", "dtwt_9": "t=10", "dtwt_14": "t=15 (MP)", "dtwt_19": "t=20", "dtwt_24": "t=25 (LP)"})
        df1.drop('Sector 1', inplace=True)

        df2 = GW_model_outputs
        df2 = df2[(df2['ibound'] == 1) & (df2['thickness'] > 0)]
        df2['initial_storage'] = df2['sat_0'] * (200 ** 2) / 1000000     # in million cubic metres
        df2['decrease_in_storage_5y'] = (df2['sat_4'] - GW_model_outputs['sat_0']) * (200 ** 2) / 1000000       # in million cubic metres
        df2['decrease_in_storage_10y'] = (df2['sat_9'] - GW_model_outputs['sat_0']) * (200 ** 2) / 1000000      # in million cubic metres
        df2['decrease_in_storage_15y'] = (df2['sat_14'] - GW_model_outputs['sat_0']) * (200 ** 2) / 1000000     # in million cubic metres
        df2['decrease_in_storage_20y'] = (df2['sat_19'] - GW_model_outputs['sat_0']) * (200 ** 2) / 1000000     # in million cubic metres
        df2['decrease_in_storage_25y'] = (df2['sat_24'] - GW_model_outputs['sat_0']) * (200 ** 2) / 1000000     # in million cubic metres
        df2 = df2.set_index('Sector_DGA')
        df2.rename(index = {1.0: "Sector 1", 2.0:"Sector 2", 3.0:"Sector 3", 4.0:"Sector 4", 5.0:"Sector 5", 6.0:"Sector 6"}, inplace = True) 
        df2 = df2.groupby('Sector_DGA').sum().reset_index().set_index('Sector_DGA')
        df_volumetric_change = df2[['decrease_in_storage_5y','decrease_in_storage_10y','decrease_in_storage_15y','decrease_in_storage_20y','decrease_in_storage_25y']]
        df_volumetric_change = df_volumetric_change.rename(index=str, columns={"decrease_in_storage_5y": "t=5 (CP)", "decrease_in_storage_10y": "t=10", "decrease_in_storage_15y": "t=15 (MP)", "decrease_in_storage_20y": "t=20", "decrease_in_storage_25y": "t=25 (LP)"})
        df_volumetric_change.drop('Sector 1', inplace=True)

        fig = plt.figure(figsize=(22,20))

        ax1 = fig.add_subplot(2,1,1)
        df1.plot(kind='bar', figsize=(10,12), colormap='viridis_r', ax=ax1)
        ax1.set_title('variación promedio de la napa freatica (m)', weight='bold', size=18)
        ax1.set_ylabel('profundidad napa (m)', weight='bold', size=15)
        ax1.set_xlabel('')
        plt.xticks(rotation=0)

        ax2 = fig.add_subplot(2,1,2)
        df_volumetric_change.plot(kind='bar', figsize=(10,12), colormap='plasma_r', ax=ax2)
        ax2.axhline(y=-250, color='red')
        ax2.axhline(y=-2500, color='red')
        plt.text(-0.45, -230, '10x volumen Embalse Lautaro', size=14)
        plt.text(-0.45, -2475, '100x volumen Embalse Lautaro', size=14)
        ax2.set_title('variación volumen embalsado (Mm3)', weight='bold', size=18)
        ax2.set_ylabel('cambio volumen c/r situación inicial (Mm3)', weight='bold', size=15)
        ax2.set_xlabel('')
        ax2.set_ylim(-3500,500)
        plt.xticks(rotation=0)

        fig.tight_layout()

        fig.savefig(self._mf(os.path.join(self.model_report_dir, '1_Evolucion_niveles_y_volumen_embalsado', '2_Scenario', 'evolucion_niveles_y_embalsado.png')), dpi=300)

        # FINAL SUMMARY RADAR PLOT

        df1['delta25'] = -(1 - df1['t=25 (LP)'] / df1['actual'])
        df1.loc[df1['delta25'] < 0, 'delta25'] = 1
        df1 = 1 - df1
        df1.loc[df1['delta25'] == 0, 'delta25'] = 1

        RADAR_SUMMARY_deltahS2_scenario = [df1['delta25'][0]][0]
        RADAR_SUMMARY_deltahS3_scenario = [df1['delta25'][1]][0]
        RADAR_SUMMARY_deltahS4_scenario = [df1['delta25'][2]][0]
        RADAR_SUMMARY_deltahS5_scenario = [df1['delta25'][3]][0]
        RADAR_SUMMARY_deltahS6_scenario = [df1['delta25'][4]][0]

        # Groundwater balance — Zonebudget heatmap

        # SUMMER

        cbc_f = os.path.join(self.model_modflow_run_dir, 'SIMCOPIAPO.cbc')
        zon = zb_array.astype(int)
        aliases = {1: 'Sector 1', 2:'Sector 2', 3: 'Sector 3', 4: 'Sector 4', 5: 'Sector 5', 6: 'Sector 6'}

        # CP_summer
        zb_CP_summer = flopy.utils.ZoneBudget(cbc_f, zon, kstpkper=(0, 18), aliases=aliases)
        zb_CP_summer = zb_CP_summer * 1000 / 86400
        zb_CP_summer_df = zb_CP_summer.get_dataframes(net = True)
        zb_CP_summer_df = zb_CP_summer_df[['Sector_2', 'Sector_3', 'Sector_4', 'Sector_5', 'Sector_6']].reset_index().drop(columns = 'totim')
        zb_CP_summer_df = zb_CP_summer_df.set_index('name')
        zb_CP_summer_df.drop(['ZONE_0', 'ZONE_999', 'CONSTANT_HEAD', 'ET'], inplace=True)
        zb_CP_summer_df = zb_CP_summer_df.reset_index().groupby('name', sort=False).sum()
        zb_CP_summer_df.rename({"STORAGE": 'Embalsado',
                      "WELLS": 'Pozos',
                      "DRAINS": 'Recuperaciones',
                      "RECHARGE": 'Recargas',
                      "Sector_1": 'Sector 1 -->', 
                      "Sector_2": 'Sector 2 -->',
                      "Sector_3": 'Sector 3 -->',
                      "Sector_4": 'Sector 4 -->',
                      "Sector_5": 'Sector 5 -->',
                      "Sector_6": 'Sector 6 -->'},
                     axis='index', inplace=True)
        zb_CP_summer_df.iloc[0] = zb_CP_summer_df.iloc[0]*-1
        zb_CP_summer_df['Sector_6']['Recuperaciones'] *= angostura_baseflow_adjustment_factor_summer
        zb_CP_summer_df['Sector_6']['Embalsado'] += zb_CP_summer_df['Sector_6']['Recuperaciones']/angostura_baseflow_adjustment_factor_summer

        # MP_summer
        zb_MP_summer = flopy.utils.ZoneBudget(cbc_f, zon, kstpkper=(0, 62), aliases=aliases)
        zb_MP_summer = zb_MP_summer * 1000 / 86400
        zb_MP_summer_df = zb_MP_summer.get_dataframes(net = True)
        zb_MP_summer_df = zb_MP_summer_df[['Sector_2', 'Sector_3', 'Sector_4', 'Sector_5', 'Sector_6']].reset_index().drop(columns = 'totim')
        zb_MP_summer_df = zb_MP_summer_df.set_index('name')
        zb_MP_summer_df.drop(['ZONE_0', 'ZONE_999', 'CONSTANT_HEAD', 'ET'], inplace=True)
        zb_MP_summer_df = zb_MP_summer_df.reset_index().groupby('name', sort=False).sum()
        zb_MP_summer_df.rename({"STORAGE": 'Embalsado',
                      "WELLS": 'Pozos',
                      "DRAINS": 'Recuperaciones',
                      "RECHARGE": 'Recargas',
                      "Sector_1": 'Sector 1 -->', 
                      "Sector_2": 'Sector 2 -->',
                      "Sector_3": 'Sector 3 -->',
                      "Sector_4": 'Sector 4 -->',
                      "Sector_5": 'Sector 5 -->',
                      "Sector_6": 'Sector 6 -->'},
                     axis='index', inplace=True)
        zb_MP_summer_df.iloc[0] = zb_MP_summer_df.iloc[0]*-1
        zb_MP_summer_df['Sector_6']['Recuperaciones'] *= angostura_baseflow_adjustment_factor_summer
        zb_MP_summer_df['Sector_6']['Embalsado'] += zb_MP_summer_df['Sector_6']['Recuperaciones']/angostura_baseflow_adjustment_factor_summer

        # LP_summer
        zb_LP_summer = flopy.utils.ZoneBudget(cbc_f, zon, kstpkper=(0, 98), aliases=aliases)
        zb_LP_summer = zb_LP_summer * 1000 / 86400
        zb_LP_summer_df = zb_LP_summer.get_dataframes(net = True)
        zb_LP_summer_df = zb_LP_summer_df[['Sector_2', 'Sector_3', 'Sector_4', 'Sector_5', 'Sector_6']].reset_index().drop(columns = 'totim')
        zb_LP_summer_df = zb_LP_summer_df.set_index('name')
        zb_LP_summer_df.drop(['ZONE_0', 'ZONE_999', 'CONSTANT_HEAD', 'ET'], inplace=True)
        zb_LP_summer_df = zb_LP_summer_df.reset_index().groupby('name', sort=False).sum()
        zb_LP_summer_df.rename({"STORAGE": 'Embalsado',
                      "WELLS": 'Pozos',
                      "DRAINS": 'Recuperaciones',
                      "RECHARGE": 'Recargas',
                      "Sector_1": 'Sector 1 -->', 
                      "Sector_2": 'Sector 2 -->',
                      "Sector_3": 'Sector 3 -->',
                      "Sector_4": 'Sector 4 -->',
                      "Sector_5": 'Sector 5 -->',
                      "Sector_6": 'Sector 6 -->'},
                     axis='index', inplace=True)
        zb_LP_summer_df.iloc[0] = zb_LP_summer_df.iloc[0]*-1
        zb_LP_summer_df['Sector_6']['Recuperaciones'] *= angostura_baseflow_adjustment_factor_summer
        zb_LP_summer_df['Sector_6']['Embalsado'] += zb_LP_summer_df['Sector_6']['Recuperaciones']/angostura_baseflow_adjustment_factor_summer

        fig = plt.figure(figsize=(10, 30))

        ax1 = fig.add_subplot(3, 1, 1)
        sns.heatmap(zb_CP_summer_df, ax=ax1, cmap='bwr_r', linewidths=1, annot=True, fmt='.0f', annot_kws={"size": 18, "weight": "bold"}, vmin=-2000, vmax=2000)
        ax1.set_title('Corto Plazo +5 años (L/s)', size=20, weight='bold')
        ax1.set_xticklabels(['Sector 2', 'Sector 3', 'Sector 4', 'Sector 5', 'Sector 6'], size=18)
        ax1.set_ylabel('')
        ax1.tick_params(axis='y', labelsize=18)
        cbar = ax1.collections[0].colorbar
        cbar.ax.tick_params(labelsize=18)

        ax2 = fig.add_subplot(3, 1, 2)
        sns.heatmap(zb_MP_summer_df, ax=ax2, cmap='bwr_r', linewidths=1, annot=True, fmt='.0f', annot_kws={"size": 18, "weight": "bold"}, vmin=-2000, vmax=2000)
        ax2.set_title('\nMediano Plazo +15 años (L/s)', size=20, weight='bold')
        ax2.set_xticklabels(['Sector 2', 'Sector 3', 'Sector 4', 'Sector 5', 'Sector 6'], size=18)
        ax2.set_ylabel('')
        ax2.tick_params(axis='y', labelsize=18)
        cbar = ax2.collections[0].colorbar
        cbar.ax.tick_params(labelsize=18)

        ax3 = fig.add_subplot(3, 1, 3)
        sns.heatmap(zb_LP_summer_df, ax=ax3, cmap='bwr_r', linewidths=1, annot=True, fmt='.0f', annot_kws={"size": 18, "weight": "bold"}, vmin=-2000, vmax=2000)
        ax3.set_title('\nLargo Plazo +25 años (L/s)', size=20, weight='bold')
        ax3.set_xticklabels(['Sector 2', 'Sector 3', 'Sector 4', 'Sector 5', 'Sector 6'], size=18)
        ax3.set_ylabel('')
        ax3.tick_params(axis='y', labelsize=18)
        cbar = ax3.collections[0].colorbar
        cbar.ax.tick_params(labelsize=18)

        fig.tight_layout()

        fig.savefig(self._mf(os.path.join(self.model_report_dir, '3_Balance_acuifero_verano', '2_Scenario', 'balance_acuifero_verano.png')), dpi=300)

        # Groundwater trends

        pozos_DGA_gdf = gpd.read_file(os.path.join(self.model_geodata_dir, 'Niveles_de_pozos_Copiapo.shp'))
        pozos_DGA_gdf = pozos_DGA_gdf[pozos_DGA_gdf.geometry.notnull()]
        pozos_DGA_gdf = pozos_DGA_gdf.to_crs(epsg=24879)[['NOMBRE', 'Sector', 'geometry']]

        sns.set(style="whitegrid")

        sectores = pozos_DGA_gdf.Sector.unique()

        fig = plt.figure(figsize=(10,25))

        for index,sector in enumerate(sectores):
            
            pozos_DGA_gdf_SX = pozos_DGA_gdf.loc[pozos_DGA_gdf['Sector'] == sector]
            
            df = pd.DataFrame()
            
            ax = fig.add_subplot(5, 1, index + 1)

            for index, row in pozos_DGA_gdf_SX.iterrows():
                
                E = row.geometry.x
                N = row.geometry.y
                nombre = row.NOMBRE
            
                row, column = ml.sr.get_rc(E,N)
                
                ts=headfile.get_ts((0, row - 1, column - 1))[:,1]
                
                df[nombre] = ts
            
                plt.plot(ts, label=nombre, linewidth=2)
            
            rate_CP = (df.iloc[10].mean() - df.iloc[0].mean()) / 5
            rate_MP = (df.iloc[30].mean() - df.iloc[10].mean()) / 10
            rate_LP = (df.iloc[49].mean() - df.iloc[30].mean()) / 10
            
            plt.title(sector, size=20, weight='bold', pad=30)
            
            plt.axvline(x=10, color='k')
            plt.axvline(x=30, color='k')
            plt.axvline(x=50, color='k')
            
            analysis_prediod_ft_size = 14
            
            plt.text(0.21, 1.01, 'CP', transform=ax.transAxes, weight='bold', size=analysis_prediod_ft_size)
            plt.text(0.59, 1.01, 'MP', transform=ax.transAxes, weight='bold', size=analysis_prediod_ft_size)
            plt.text(0.96, 1.01, 'LP', transform=ax.transAxes, weight='bold', size=analysis_prediod_ft_size)
            
            rate_ft_size = 16
            
            if rate_CP >= 0:
                plt.text(0.07, 1.01, '+' + str(round(rate_CP,1)) + ' m/año', transform=ax.transAxes, color='blue', size=rate_ft_size)
            else:
                plt.text(0.07, 1.01, str(round(rate_CP,1)) + ' m/año', transform=ax.transAxes, color='red', size=rate_ft_size)
            
            if rate_MP >= 0:
                plt.text(0.37, 1.01, '+' + str(round(rate_MP,1)) + ' m/año', transform=ax.transAxes, color='blue', size=rate_ft_size)
            else:
                plt.text(0.37, 1.01, str(round(rate_MP,1)) + ' m/año', transform=ax.transAxes, color='red', size=rate_ft_size)

            if rate_LP >= 0:
                plt.text(0.75, 1.01, '+' + str(round(rate_MP,1)) + ' m/año', transform=ax.transAxes, color='blue', size=rate_ft_size)
            else:
                plt.text(0.75, 1.01, str(round(rate_MP,1)) + ' m/año', transform=ax.transAxes, color='red', size=rate_ft_size)
            
            plt.xticks(np.arange(0, 51, step=10), ('2019','2024','2029','2033','2038','2043'))
            plt.xlabel('')
            plt.ylabel('nivel (msnm)', weight='bold')
            plt.legend(loc=3, prop={'size': 6})
            plt.ylabel('nivel (msnm)', weight='bold', size=16)
            plt.legend(loc=3, prop={'size': 10})

        fig.tight_layout()

        fig.savefig(self._mf(os.path.join(self.model_report_dir, '2_Niveles_freaticos_pozos_DGA', '2_Scenario', 'niveles_freaticos_pozos_DGA.png')), dpi=300)

        # Operacion Embalse Lautaro

        fig = plt.figure(figsize=(18,10))

        gs = GridSpec(5, 3, figure=fig)
        ax1 = fig.add_subplot(gs[0,0])
        ax2 = fig.add_subplot(gs[1,0])
        ax3 = fig.add_subplot(gs[2,0])
        ax4 = fig.add_subplot(gs[3,0])
        ax5 = fig.add_subplot(gs[4,0])
        ax6 = fig.add_subplot(gs[:,1])
        ax7 = fig.add_subplot(gs[:,2])

        data = SWMODEL_out_df[['Q_Pastillo',
                               'Q_LautaroControlada',
                               'Q_Vertedero',
                               'Q_InfiltracionLautaro',
                               'P_full']].head(300)

        data = data.tshift(12*27, freq='M')

        data.Q_Pastillo.plot(ax=ax1, label='Q Pastillos')
        data.Q_LautaroControlada.plot(ax=ax2, label='Q descarga Lautaro')
        data.Q_Vertedero.plot(ax=ax3, label='Q vertedero', color='red')
        data.Q_InfiltracionLautaro.plot(ax=ax4, label='Q infiltracion fondo')
        data.P_full.plot(ax=ax5, label='% capacidad')

        if ST_lautaro2:
            Lautaro2_axvline_color = 'r'
            ax1.axvline('2024-1-1', color=Lautaro2_axvline_color, linestyle=':', linewidth=2)
            ax2.axvline('2024-1-1', color=Lautaro2_axvline_color, linestyle=':', linewidth=2)
            ax3.axvline('2024-1-1', color=Lautaro2_axvline_color, linestyle=':', linewidth=2)
            ax4.axvline('2024-1-1', color=Lautaro2_axvline_color, linestyle=':', linewidth=2)
            ax5.axvline('2024-1-1', color=Lautaro2_axvline_color, linestyle=':', linewidth=2)

        ax1.set_title('Principales entradas, salidas y acopio embalse', size=14, weight='bold');

        ax5.axhline(y=0.5, color='r', linestyle='-')

        ax1.set_xlabel('');
        ax2.set_xlabel('');
        ax3.set_xlabel('');
        ax4.set_xlabel('');
        ax5.set_xlabel('');

        ax1.set_ylabel('L/s', weight='bold');
        ax2.set_ylabel('L/s', weight='bold');
        ax3.set_ylabel('L/s', weight='bold');
        ax4.set_ylabel('L/s', weight='bold');
        ax5.set_ylabel('%', weight='bold');
        ax5.set_yticks([0, 0.25, 0.5, 0.75, 1]);
        ax5.set_yticklabels(('0', '25', '50', '75', '100'));

        ax1.legend()
        ax2.legend()
        ax3.legend()
        ax4.legend()
        ax5.legend()

        ax5.fill_between(data.P_full.index, data.P_full.values, 0, color='blue')  

        data = SWMODEL_out_df['Q_Vertedero'].head(300)
        data = (data > 0).astype(int)

        xticklabels=['SEP', 'OCT', 'NOV', 'DIC', 'ENE', 'FEB', 'MAR', 'ABR', 'MAY', 'JUN', 'JUL', 'AGO' ]
        yticklabels=['2018/19','2019/20', '2020/21','2021/22','2022/23','2023/24','2024/25','2025/26','2026/27','2027/28','2028/29','2029/30','2030/31','2031/32','2032/33','2033/34','2034/35','2035/36','2036/37','2037/38','2038/39','2039/40','2040/41','2041/42','2042/43']

        n_months_with_spillage = int(data.sum()/300)

        sns.heatmap(data.values.reshape(-1,12), linewidths=0.3, linecolor='k', cmap='bwr', cbar=False, xticklabels=xticklabels, yticklabels=yticklabels, ax=ax6)

        ax6.set_title('meses vertedero sobrepasado\n| rojo=rebalse | azul=sin rebalse |', size=18, weight='bold');

        data = SWMODEL_out_df['P_full'].head(300)
        data = (data > 0.5).astype(int)

        n_months_over50pctfull = int(data.sum()/300)

        sns.heatmap(data.values.reshape(-1,12), linewidths=0.3, linecolor='k', cmap='bwr_r', cbar=False, xticklabels=xticklabels, yticklabels=yticklabels, ax=ax7)

        ax7.set_title('meses embalse con >50% capacidad\n| azul=sobre | rojo=bajo |', size=18, weight='bold');

        fig.tight_layout()

        fig.savefig(self._mf(os.path.join(self.model_report_dir, '5_Embalse_Lautaro', '2_Scenario', 'embalse_lautaro.png')), dpi=300)

        # FINAL SUMMARY RADAR PLOT

        RADAR_SUMMARY_LautaroVertedero_scenario = (SWMODEL_out_df['Q_Vertedero'].head(300).values == 0).sum()/300
        RADAR_SUMMARY_Lautaro50_scenario = (SWMODEL_out_df['P_full'].head(300).values >= 0.5).sum()/300

        # JVRC

        fig = plt.figure(figsize=(10,10))

        gs = GridSpec(4, 2, figure=fig)
        ax1 = fig.add_subplot(gs[0,0])
        ax2 = fig.add_subplot(gs[0,1])
        ax3 = fig.add_subplot(gs[1,0])
        ax4 = fig.add_subplot(gs[1,1])
        ax5 = fig.add_subplot(gs[2,0])
        ax6 = fig.add_subplot(gs[2,1])
        ax7 = fig.add_subplot(gs[3,0])
        ax8 = fig.add_subplot(gs[3,1])

        data = SWMODEL_out_df[['Satisfaccion_dda_C_D1',
                               'Satisfaccion_dda_C_D2',
                               'Satisfaccion_dda_C_D3',
                               'Satisfaccion_dda_C_D4',
                               'Satisfaccion_dda_C_D5',
                               'Satisfaccion_dda_C_D6',
                               'Satisfaccion_dda_C_D7',
                               'Satisfaccion_dda_C_D89']]

        data = data.tshift(12*27, freq='M').head(300)

        data.Satisfaccion_dda_C_D1.plot(ax=ax1, style='o', label='Distrito 1')
        data.Satisfaccion_dda_C_D2.plot(ax=ax2, style='o', label='Distrito 2')
        data.Satisfaccion_dda_C_D3.plot(ax=ax3, style='o', label='Distrito 3')
        data.Satisfaccion_dda_C_D4.plot(ax=ax4, style='o', label='Distrito 4')
        data.Satisfaccion_dda_C_D5.plot(ax=ax5, style='o', label='Distrito 5')
        data.Satisfaccion_dda_C_D6.plot(ax=ax6, style='o', label='Distrito 6')
        data.Satisfaccion_dda_C_D7.plot(ax=ax7, style='o', label='Distrito 7')
        data.Satisfaccion_dda_C_D89.plot(ax=ax8, style='o', label='Distrito 89')


        if ST_lautaro2:
            Lautaro2_axvline_color = 'r'
            ax1.axvline('2024-1-1', color=Lautaro2_axvline_color, linestyle=':', linewidth=2)
            ax2.axvline('2024-1-1', color=Lautaro2_axvline_color, linestyle=':', linewidth=2)
            ax3.axvline('2024-1-1', color=Lautaro2_axvline_color, linestyle=':', linewidth=2)
            ax4.axvline('2024-1-1', color=Lautaro2_axvline_color, linestyle=':', linewidth=2)
            ax5.axvline('2024-1-1', color=Lautaro2_axvline_color, linestyle=':', linewidth=2)
            ax6.axvline('2024-1-1', color=Lautaro2_axvline_color, linestyle=':', linewidth=2)
            ax7.axvline('2024-1-1', color=Lautaro2_axvline_color, linestyle=':', linewidth=2)
            ax8.axvline('2024-1-1', color=Lautaro2_axvline_color, linestyle=':', linewidth=2)


        ax1.legend(loc=3, prop={'size': 10});
        ax2.legend(loc=3, prop={'size': 10});
        ax3.legend(loc=3, prop={'size': 10});
        ax4.legend(loc=3, prop={'size': 10});
        ax5.legend(loc=3, prop={'size': 10});
        ax6.legend(loc=3, prop={'size': 10});
        ax7.legend(loc=3, prop={'size': 10});
        ax8.legend(loc=3, prop={'size': 10});

        ax1.set_yticks([0,0.25,0.5,0.75,1])
        ax2.set_yticks([0,0.25,0.5,0.75,1])
        ax3.set_yticks([0,0.25,0.5,0.75,1])
        ax4.set_yticks([0,0.25,0.5,0.75,1])
        ax5.set_yticks([0,0.25,0.5,0.75,1])
        ax6.set_yticks([0,0.25,0.5,0.75,1])
        ax7.set_yticks([0,0.25,0.5,0.75,1])
        ax8.set_yticks([0,0.25,0.5,0.75,1])

        ax1.set_xlabel('')
        ax2.set_xlabel('')
        ax3.set_xlabel('')
        ax4.set_xlabel('')
        ax5.set_xlabel('')
        ax6.set_xlabel('')
        ax7.set_xlabel('')
        ax8.set_xlabel('')

        ax1.set_yticklabels(['0%','25%','50%','75%','100%'])
        ax2.set_yticklabels(['0%','25%','50%','75%','100%'])
        ax3.set_yticklabels(['0%','25%','50%','75%','100%'])
        ax4.set_yticklabels(['0%','25%','50%','75%','100%'])
        ax5.set_yticklabels(['0%','25%','50%','75%','100%'])
        ax6.set_yticklabels(['0%','25%','50%','75%','100%'])
        ax7.set_yticklabels(['0%','25%','50%','75%','100%'])
        ax8.set_yticklabels(['0%','25%','50%','75%','100%'])

        legend_properties = {'size':16}

        ax1.legend(prop=legend_properties, loc=1)
        ax2.legend(prop=legend_properties, loc=1)
        ax3.legend(prop=legend_properties, loc=1)
        ax4.legend(prop=legend_properties, loc=1)
        ax5.legend(prop=legend_properties, loc=1)
        ax6.legend(prop=legend_properties, loc=1)
        ax7.legend(prop=legend_properties, loc=1)
        ax8.legend(prop=legend_properties, loc=1)

        fig.suptitle('% satisfaccion demanda mensual distritos', weight='bold', size=18)

        fig.tight_layout()

        fig.subplots_adjust(top=0.95)

        fig.savefig(self._mf(os.path.join(self.model_report_dir, '6_JVRC','2_Scenario','satisfaccion_dda_ts.png')), dpi=300)

        # Heatmaps

        fig = plt.figure(figsize=(10,10))

        gs = GridSpec(1, 2, figure=fig)
        ax1 = fig.add_subplot(gs[0,0])
        ax2 = fig.add_subplot(gs[0,1])

        data = SWMODEL_out_df[['Satisfaccion_dda_C_D1',
                               'Satisfaccion_dda_C_D2',
                               'Satisfaccion_dda_C_D3',
                               'Satisfaccion_dda_C_D4',
                               'Satisfaccion_dda_C_D5',
                               'Satisfaccion_dda_C_D6',
                               'Satisfaccion_dda_C_D7',
                               'Satisfaccion_dda_C_D89']]

        data = data.tshift(12*27, freq='M').head(300)

        sns.heatmap(data, cmap='coolwarm_r', yticklabels = 12, cbar=True, cbar_kws={"aspect": 40}, ax=ax1)
        ax1.set_xticklabels(['Distrito 1', 'Distrito 2', 'Distrito 3', 'Distrito 4', 'Distrito 5', 'Distrito 6', 'Distrito 7', 'Distritos 89']);
        ax1.set_yticklabels(data.iloc[::12, :].index.strftime("%Y"));
        ax1.set_title('satisfaccion demanda\npredios solo canal', size=16, weight='bold');
        ax1.set_ylabel('');
        ax1.tick_params(axis='y', labelsize=12)
        cbar = ax1.collections[0].colorbar
        cbar.ax.tick_params(labelsize=12)

        data = SWMODEL_out_df[['Q_dda_M_a_P_D1',
                               'Q_dda_M_a_P_D2',
                               'Q_dda_M_a_P_D3',
                               'Q_dda_M_a_P_D4',
                               'Q_dda_M_a_P_D5',
                               'Q_dda_M_a_P_D6',
                               'Q_dda_M_a_P_D7',
                               'Q_dda_M_a_P_D89']]

        sns.heatmap(data, cmap='inferno', yticklabels = 12, cbar=True, cbar_kws={"aspect": 40}, ax=ax2, vmin=0, vmax=300)
        ax2.set_xticklabels(['Distrito 1', 'Distrito 2', 'Distrito 3', 'Distrito 4', 'Distrito 5', 'Distrito 6', 'Distrito 7', 'Distritos 89']);
        ax2.set_yticklabels(data.iloc[::12, :].index.strftime("%Y"));
        ax2.set_title('bombeo adicional por\ninsatisfacción demanda (L/s)', size=16, weight='bold');
        ax2.set_ylabel('');
        ax2.tick_params(axis='both', labelsize=12)
        cbar = ax2.collections[0].colorbar
        cbar.ax.tick_params(labelsize=12)

        fig.tight_layout()

        fig.savefig(self._mf(os.path.join(self.model_report_dir, '6_JVRC', '2_Scenario', 'satisfaccion_dda_heatmap.png')), dpi=300)

        # Radar Plot

        data = SWMODEL_out_df[['Satisfaccion_dda_C_D1',
                       'Satisfaccion_dda_C_D2',
                       'Satisfaccion_dda_C_D3',
                       'Satisfaccion_dda_C_D4',
                       'Satisfaccion_dda_C_D5',
                       'Satisfaccion_dda_C_D6',
                       'Satisfaccion_dda_C_D7',
                       'Satisfaccion_dda_C_D89']]

        data = data.tshift(12*27, freq='M').head(300)

        fig = plt.figure(figsize=(6,6))

        # number of variables
        N = 8

        # We are going to plot the first line of the data frame.
        # But we need to repeat the first value to close the circular graph:
        values=(data.sum()/300).values.tolist()
        values += values[:1]

        # What will be the angle of each axis in the plot? (we divide the plot / number of variable)
        angles = [n / float(N) * 2 * pi for n in range(N)]
        angles += angles[:1]

        # Initialise the spider plot
        ax = plt.subplot(111, polar=True)

        # Draw one axe per variable + add labels labels yet
        plt.xticks(angles[:-1], ['D1', 'D2', 'D3', 'D4', 'D5', 'D6', 'D7', 'D89'], size=16, weight='bold')

        # Draw ylabels
        ax.set_rlabel_position(0)
        plt.yticks([0, 0.2,0.4,0.6, 0.8, 1], ["0%", "20%","40%","60%","80%", "100%"], color="k", size=14)
        plt.ylim(0,1)

        # Plot data
        ax.plot(angles, values, linewidth=1, linestyle='solid')

        # Fill area
        ax.fill(angles, values, 'b', alpha=0.2)

        ax.set_title('satisfaccion demandas superficiales (media)', size=16, pad=20, weight='bold');

        fig.tight_layout()

        fig.savefig(self._mf(os.path.join(self.model_report_dir, '6_JVRC', '2_Scenario', 'satisfaccion_dda_radar.png')), dpi=300)

        # Rio Copiapo

        fig = plt.figure(figsize=(20,12))

        gs = GridSpec(5, 3, figure=fig)
        ax1 = fig.add_subplot(gs[0,0])
        ax2 = fig.add_subplot(gs[1,0])
        ax3 = fig.add_subplot(gs[2,0])
        ax4 = fig.add_subplot(gs[3,0])
        ax5 = fig.add_subplot(gs[4,0])
        ax6 = fig.add_subplot(gs[:,1])
        ax7 = fig.add_subplot(gs[:,2])

        data = SWMODEL_out_df[['Q_LautaroControlada',
                               'Q_LaPuerta',
                               'Q_rio_Nantoco',
                               'Q_rio_CCopiapo',
                               'Q_rio_Angostura']]

        data = data.tshift(12*27, freq='M').head(300)

        data.Q_LautaroControlada.plot(ax=ax1, label='Q Lautaro')
        data.Q_LaPuerta.plot(ax=ax2, label='Q La Puerta')
        data.Q_rio_Nantoco.plot(ax=ax3, label='Q Mal Paso')
        data.Q_rio_CCopiapo.plot(ax=ax4, label='Q Ciudad')
        data.Q_rio_Angostura.plot(ax=ax5, label='Q Angostura')


        if ST_lautaro2:
            Lautaro2_axvline_color = 'r'
            ax1.axvline('2024-1-1', color=Lautaro2_axvline_color, linestyle=':', linewidth=2)
            ax2.axvline('2024-1-1', color=Lautaro2_axvline_color, linestyle=':', linewidth=2)
            ax3.axvline('2024-1-1', color=Lautaro2_axvline_color, linestyle=':', linewidth=2)
            ax4.axvline('2024-1-1', color=Lautaro2_axvline_color, linestyle=':', linewidth=2)
            ax5.axvline('2024-1-1', color=Lautaro2_axvline_color, linestyle=':', linewidth=2)

        ax1.set_title('series hidrologicas\nen estaciones fluviometricas DGA', weight='bold', size=16)

        ax1.set_ylabel('L/s', weight='bold')
        ax2.set_ylabel('L/s', weight='bold')
        ax3.set_ylabel('L/s', weight='bold')
        ax4.set_ylabel('L/s', weight='bold')
        ax5.set_ylabel('L/s', weight='bold')

        ax1.set_ylabel('L/s', weight='bold', size=16)
        ax2.set_ylabel('L/s', weight='bold', size=16)
        ax3.set_ylabel('L/s', weight='bold', size=16)
        ax4.set_ylabel('L/s', weight='bold', size=16)
        ax5.set_ylabel('L/s', weight='bold', size=16)

        ax1.set_xlabel('')
        ax2.set_xlabel('')
        ax3.set_xlabel('')
        ax4.set_xlabel('')
        ax5.set_xlabel('')

        legend_properties = {'size':16}

        ax1.legend(prop=legend_properties)
        ax2.legend(prop=legend_properties)
        ax3.legend(prop=legend_properties)
        ax4.legend(prop=legend_properties)
        ax5.legend(prop=legend_properties)

        medians = data.median().values
        median_labels = [str(int(s)) for s in medians]

        sns.boxplot(data=data, showfliers=False, palette='Blues', ax=ax6)
        ax6.set_title('estadisticas\nen estaciones fluviometricas DGA', size=16, weight='bold')
        ax6.set_ylabel('[L/s]', weight='bold', size=16);
        ax6.set_xticklabels(['Lautaro', 'La Puerta', 'Mal Paso', 'Ciudad', 'Angostura'], size=16)

        pos = range(len(medians))
        for tick,label in zip(pos,ax.get_xticklabels()):
            ax6.text(pos[tick], medians[tick] + 15, median_labels[tick], 
                    horizontalalignment='center', size=18, color='k', weight='bold')

        data = SWMODEL_out_df[['RCH_rio_S2',
                               'RCH_rio_S3',
                               'RCH_rio_S4',
                               'RCH_rio_S5',
                               'RCH_rio_S6',]]

        medians = data.median().values
        median_labels = [str(int(s)) for s in medians]

        sns.boxplot(data=data, showfliers=False, palette='Purples', ax=ax7)
        ax7.set_title('recarga natural\nen Rio Copiapo', size=16, weight='bold')
        ax7.set_ylabel('[L/s]', weight='bold');
        ax7.set_xticklabels(['Sector 2', 'Sector 3', 'Sector 4', 'Sector 5', 'Sector 6']);

        pos = range(len(medians))
        for tick,label in zip(pos,ax.get_xticklabels()):
            ax7.text(pos[tick], medians[tick] + 7, median_labels[tick], 
                    horizontalalignment='center', size=18, color='k', weight='bold')

        fig.tight_layout()

        fig.savefig(self._mf(os.path.join(self.model_report_dir, '7_Rio_Copiapo', '2_Scenario', 'rio_copiapo.png')), dpi=300)

        # Caudal Social y Ecologico

        S6_zb_df = zb_df.Sector_6.reset_index()
        S6_zb_df = S6_zb_df.loc[((S6_zb_df['name'] == 'TO_DRAINS'))]
        #mean = S6_zb_df.groupby('datetime').sum().mean().values[0]

        anf_df=S6_zb_df.groupby('datetime').sum().reset_index()
        anf_df.loc[anf_df['datetime'].dt.month == 6, ['Sector_6']] *= angostura_baseflow_adjustment_factor_summer
        anf_df.loc[anf_df['datetime'].dt.month == 12, ['Sector_6']] *= angostura_baseflow_adjustment_factor_winter
        anf_df.set_index('datetime', inplace=True)
        a=anf_df.resample('M').pad().head(300).tshift(6, freq='M')

        aporte_subterraneo_df = a
        aporte_superficial_df = SWMODEL_out_df['Q_rio_Angostura'].tshift(12*27, freq='M').head(300)
        aporte_total = (a.Sector_6 + aporte_superficial_df)

        df_Angostura = pd.DataFrame()
        df_Angostura['aporte_superficial'] = aporte_superficial_df
        df_Angostura['aporte_subterraneo'] = aporte_subterraneo_df
        df_Angostura['total'] = aporte_total

        plt.figure(figsize=(12,8))

        plt.plot(df_Angostura.aporte_subterraneo + df_Angostura.aporte_superficial, label='superficial + subterraneo')
        plt.title('caudal humedal sector Angostura (aportes superficiales y subterraneos)', weight='bold', size=16)
        plt.ylabel('L/s', weight='bold');
        plt.legend()
        plt.ylim(0,300)
        plt.axhline(y=caudal_ecologico_historico, color='g', linestyle='-')

        if ST_lautaro2:
            Lautaro2_axvline_color = 'r'
            plt.axvline('2024-1-1', color=Lautaro2_axvline_color, linestyle=':', linewidth=2)

        plt.text('2036-6-1', caudal_ecologico_historico+10, 'caudal ecologico historico\n     (minimo) pre-1985', color='k', bbox=dict(boxstyle="round",
                           ec=(1., 0.5, 0.5),
                           fc=(1., 0.8, 0.8),))

        plt.savefig(self._mf(os.path.join(self.model_report_dir, '8_Q_Social_y_Ambiental', '2_Scenario', 'Q_humedal.png')), dpi=300, bbox_inches='tight')

        caudal_social_max = 5000

        data = SWMODEL_out_df[['Q_rio_CCopiapo',
                               'Q_rio_Angostura']].head(288)

        df_Angostura = df_Angostura.head(288)

        data.index = df_Angostura.index

        data['Q_rio_Angostura'] = df_Angostura['total']

        data.loc[data['Q_rio_CCopiapo'] == 0, 'Q_rio_CCopiapo_case'] = 0
        data.loc[((data['Q_rio_CCopiapo'] > 0) & (data['Q_rio_CCopiapo'] < caudal_social_max)), 'Q_rio_CCopiapo_case'] = 1
        data.loc[data['Q_rio_CCopiapo'] > caudal_social_max, 'Q_rio_CCopiapo_case'] = 2

        data.loc[((data['Q_rio_Angostura'] > 0) & (data['Q_rio_Angostura'] < caudal_ecologico_historico)), 'Q_rio_Angostura_case'] = 0
        data.loc[data['Q_rio_Angostura'] >= caudal_ecologico_historico, 'Q_rio_Angostura_case'] = 1
        data['Q_rio_Angostura_case'].fillna(value=0, inplace=True)

        fig = plt.figure(figsize=(12,8))

        gs = GridSpec(1, 2, figure=fig)
        ax1 = fig.add_subplot(gs[0,0])
        ax2 = fig.add_subplot(gs[0,1])

        xticklabels=['SEP', 'OCT', 'NOV', 'DIC', 'ENE', 'FEB', 'MAR', 'ABR', 'MAY', 'JUN', 'JUL', 'AGO' ]
        yticklabels=['2018/19','2019/20', '2020/21','2021/22','2022/23','2023/24','2024/25','2025/26','2026/27','2027/28','2028/29','2029/30','2030/31','2031/32','2032/33','2033/34','2034/35','2035/36','2036/37','2037/38','2038/39','2039/40','2040/41','2041/42']

        mypalette_1 = ['#EFEFEF', '#262DF3', '#CD853F']
        mypalette_2 = ['#FD4118', '#DAFD82']

        sns.heatmap(data['Q_rio_CCopiapo_case'].values.reshape(-1,12), linewidths=0.3, linecolor='k', cmap=mypalette_1, cbar=True, ax=ax1, xticklabels=xticklabels, yticklabels=yticklabels, vmin=0, vmax=2, cbar_kws={"orientation": "horizontal", "pad": 0.05, "ticks":[0.333, 0.666, 1, 1.333, 1.666]})
        sns.heatmap(data['Q_rio_Angostura_case'].values.reshape(-1,12), linewidths=0.3, linecolor='k', cmap=mypalette_2, cbar=True, ax=ax2, xticklabels=xticklabels, yticklabels=yticklabels, cbar_kws={"orientation": "horizontal", "pad": 0.05, "ticks":[0.25,0.5,0.75]})

        ax1.set_title('caudal urbano (Q Rio Copiapo en Ciudad)', weight='bold', size=14);
        ax2.set_title('caudal humedal (Q Rio Copiapo en Angostura)', weight='bold', size=14);

        cbar_ax1 = ax1.collections[0].colorbar
        cbar_ax2 = ax2.collections[0].colorbar

        cbar_ax1.ax.set_xticklabels(['seco', '0 L/s', 'aceptable', str(caudal_social_max) +' L/s\n(Pr=5)', 'parque\ninundado']) 
        cbar_ax1.ax.tick_params(labelsize=12)

        cbar_ax2.ax.set_xticklabels(['bajo minimo\\nhistorico', str(caudal_ecologico_historico) + ' L/s', 'sobre minimo historico']) 
        cbar_ax2.ax.tick_params(labelsize=12)

        fig.tight_layout()

        fig.savefig(self._mf(os.path.join(self.model_report_dir, '8_Q_Social_y_Ambiental', '2_Scenario', 'Q_social_y_ambiental.png')), dpi=300)

        ### FINAL SUMMARY RADAR PLOT ###
        RADAR_SUMMARY_caudal_social_scenario = (data['Q_rio_CCopiapo_case'].values == 1).sum()/288
        RADAR_SUMMARY_caudal_ecologico_scenario = (data['Q_rio_Angostura'].values >=caudal_ecologico_historico).sum()/288

        # Indicadores economicos

        df = GW_model_outputs

        df = df[(df['ibound'] == 1) & (df['thickness'] > 0)]
        df = df[['Sector_DGA','dtwt_0','dtwt_4','dtwt_9','dtwt_14','dtwt_19','dtwt_24']]
        df = df.groupby('Sector_DGA').mean().reset_index().set_index('Sector_DGA')
        df.rename(index = {1.0: "Sector 1", 2.0:"Sector 2", 3.0:"Sector 3", 4.0:"Sector 4", 5.0:"Sector 5", 6.0:"Sector 6"}, inplace = True)
        df = df.rename(index=str, columns={"dtwt_0": "t=0", "dtwt_4": "t=5", "dtwt_9": "t=10", "dtwt_14": "t=15", "dtwt_19": "t=20", "dtwt_24": "t=25"})
        df.drop('Sector 1', inplace=True)

        df['t=0'] = df['t=0'] * (2.725 / 0.7 / 0.95 / 0.9 * 80)
        df['t=5'] = df['t=5'] * (2.725 / 0.7 / 0.95 / 0.9 * 80)
        df['t=10'] = df['t=10'] * (2.725 / 0.7 / 0.95 / 0.9 * 80)
        df['t=15'] = df['t=15'] * (2.725 / 0.7 / 0.95 / 0.9 * 80)
        df['t=20'] = df['t=20'] * (2.725 / 0.7 / 0.95 / 0.9 * 80)
        df['t=25'] = df['t=25'] * (2.725 / 0.7 / 0.95 / 0.9 * 80)

        fig = plt.figure(figsize=(20,6))

        gs = GridSpec(1, 2, figure=fig)
        ax1 = fig.add_subplot(gs[0,0])
        ax2 = fig.add_subplot(gs[0,1])

        df.plot(kind='bar', colormap='plasma_r', ax=ax1)
        ax1.set_title('Costo de Bombeo (referencial)', weight='bold', size=16)
        ax1.set_ylabel('costo (CLP/1000m3)', weight='bold', size=16)
        ax1.set_xlabel('')
        formatter = ticker.FormatStrFormatter('$%1.0f')
        ax1.yaxis.set_major_formatter(formatter)
        ax1.set_xticklabels(['Sector 2', 'Sector 3', 'Sector 4', 'Sector 5', 'Sector 6'], rotation = 0, size=14)

        DW_MODEL_df['Desaladora'].head(25).plot(kind='bar', ax=ax2)
        ax2.set_title('Operacion Planta Desaladora Estatal', weight = 'bold', size=16)
        ax2.set_ylabel('Caudal agua desalada [L/s]', weight = 'bold', size=16)
        ax2.set_xlabel('')
        ax2.set_xticklabels(DW_MODEL_df['Desaladora'].head(25).index.year.values);
        plt.xticks(rotation='horizontal', size=10)

        fig.tight_layout()

        fig.savefig(self._mf(os.path.join(self.model_report_dir, '9_Economico', '2_Scenario', 'indicadores_economicos.png')), dpi=300, bbox_inches='tight')

        # Indicadores resumen

        # number of variable
        N = 9

        # What will be the angle of each axis in the plot? (we divide the plot / number of variable)
        angles = [n / float(N) * 2 * pi for n in range(N)]
        angles += angles[:1]

        # Initialise the spider plot
        fig = plt.figure(figsize=(10, 10))
        ax = plt.subplot(111, polar=True)

        # If you want the first axis to be on top:
        ax.set_theta_offset(pi / 2)
        ax.set_theta_direction(-1)

        # Draw ylabels
        ax.set_rlabel_position(0)
        plt.yticks([0, 0.2,0.4,0.6, 0.8, 1], ["0%", "20%","40%","60%","80%", "100%"], color="grey", size=12)
        plt.ylim(0,1)

        # Plot each individual = each line of the data

        # base
        values_base=[0.17708333333333334,
         0.4618055555555556,
         0.21333333333333335,
         1.0,
         0.8321432975323426,
         0.6514795120813572,
         0.21208550476303123,
         0.731660032324843,
         0]

        values_base += values_base[:1]
        ax.plot(angles, values_base, linewidth=2, linestyle='solid', label="escenario base")
        ax.fill(angles, values_base, 'b', alpha=0.1)

        # Draw one axe per variable + add labels labels yet
        plt.xticks(angles[:-1], ['caudal\\nurbano\\nentre 0 y 5000L/s', 
                                 'caudal\\necologico\\nsobre\\n>%s L/s'%caudal_ecologico_historico, 
                                 'volumen\\nLautaro\n>50%', 
                                 '% volumen\\ninicial\\nAcuifero\\nS2', 
                                 '% volumen\\ninicial\\nAcuifero\\nS3', 
                                 '% volumen\\ninicial\\nAcuifero\\nS4', 
                                 '% volumen\\ninicial\\nAcuifero\\nS5', 
                                 '% volumen\\ninicial\\nAcuifero\\nS6', 
                                 '% cumplimiento\\nprorratas'], size=16, weight='bold')

        ax.tick_params(axis='both', which='major', pad=25)

        # scenario
        values_scenario=[RADAR_SUMMARY_caudal_social_scenario, 
                     RADAR_SUMMARY_caudal_ecologico_scenario, 
                     RADAR_SUMMARY_Lautaro50_scenario,
                     RADAR_SUMMARY_deltahS2_scenario,
                     RADAR_SUMMARY_deltahS3_scenario,
                     RADAR_SUMMARY_deltahS4_scenario,
                     RADAR_SUMMARY_deltahS5_scenario,
                     RADAR_SUMMARY_deltahS6_scenario,
                     compliance]
        values_scenario += values_scenario[:1]
        ax.plot(angles, values_scenario, linewidth=2, linestyle='solid', label="escenario simulado")
        ax.fill(angles, values_scenario, 'r', alpha=0.1)

        # Add legend
        plt.legend(bbox_to_anchor=(0.5, 0.9, 0.20, 0.4), fontsize=18)

        fig.savefig(self._mf(os.path.join(self.model_report_dir, '10_Radar_Resumen', 'radar_resumen.png')), dpi=300, bbox_inches='tight')

        # Cuenca (scenario - baseline) map + radar inset

        from matplotlib.cbook import get_sample_data

        fig, ax = plt.subplots(figsize=(15, 15))

        plt.title('diferencia napa freatica (metros) situacion base y simulada\n azul=mejora | rojo=empeora', size=22, weight='bold')

        im = ax.imshow(heads_scenario_improvement, vmin=-50, vmax=50, cmap='seismic_r')
        cax = fig.add_axes([0.8, 0.37, 0.01, 0.5])

        fig.colorbar(im, cax=cax)

        axin = fig.add_axes([0.1, -0.10, 0.55, 0.55])

        radar_im = plt.imread(os.path.join(self.model_report_data_dir, '10_Radar_Resumen', 'radar_resumen.png'))

        axin.imshow(radar_im)

        ax.grid(False)
        ax.set_xticks([]);
        ax.set_yticks([]);

        axin.grid(False)
        axin.set_xticks([]);
        axin.set_yticks([]);
        axin.set_title('resumen indicadores', size=22, weight='bold');

        fig.tight_layout()

        fig.savefig(self._mf(os.path.join(self.model_report_dir, '10_Radar_Resumen', 'radar_y_mapa_resumen.png' )), dpi=300, bbox_inches='tight')

        # Build report using python-docx workflow

        now = dt.datetime.now()

        document = Document(self._mf(os.path.join(self.model_report_data_dir, '_Tapa_informe', 'cover_1.docx')))

        tables = document.tables

        tables[0].rows[0].cells[0].text = str('creado el %s' %now.strftime("%d-%m-%Y a las %H:%M"))

        document.add_section(WD_SECTION.NEW_PAGE)

        document.save(self._mf(os.path.join(self.model_report_dir, '_Tapa_informe', 'cover_out.docx')))

        # Page 1 — Configuracion escenario base y escenario simulado

        # ***** 

        # Page 2 — Resumen situacion cuenca

        now = dt.datetime.now()

        document = Document(os.path.join(self.model_report_data_dir, '10_Radar_Resumen', 'report_template.docx'))

        style = document.styles['Normal']
        font = style.font
        font.name = 'Avenir Light'
        font.size = Pt(10)

        tables = document.tables

        p = tables[0].rows[0].cells[0].add_paragraph()
        r = p.add_run()
        r.add_picture(self._mf(os.path.join(self.model_report_dir, '10_Radar_Resumen', 'radar_y_mapa_resumen.png')), width = Cm(20))

        # section = document.sections[0]
        # footer = section.footer
        # footer.add_paragraph('Reporte Simulacion SIMCOPIAPO | creado el %s' %now.strftime("%d-%m-%Y a las %H:%M"))

        document.add_section(WD_SECTION.NEW_PAGE)

        document.save(self._mf(os.path.join(self.model_report_dir, '10_Radar_Resumen', 'report_out.docx')))

        # Page 3 — Evolucion niveles freaticos y cambio en volumen embalsado

        now = dt.datetime.now()

        document = Document(os.path.join(self.model_report_data_dir, '1_Evolucion_niveles_y_volumen_embalsado', 'report_template.docx'))

        style = document.styles['Normal']
        font = style.font
        font.name = 'Avenir Light'
        font.size = Pt(10)

        tables = document.tables

        p = tables[0].rows[0].cells[0].add_paragraph()
        r = p.add_run()
        r.add_picture(os.path.join(self.model_report_data_dir, '1_Evolucion_niveles_y_volumen_embalsado', '1_Baseline', 'evolucion_niveles_y_embalsado.png'), width = Cm(12.5))

        # ------> REPLACE WITH SCENARIO FIGURE <------
        p = tables[0].rows[0].cells[1].add_paragraph()
        r = p.add_run()
        r.add_picture(self._mf(os.path.join(self.model_report_dir, '1_Evolucion_niveles_y_volumen_embalsado', '2_Scenario', 'evolucion_niveles_y_embalsado.png')), width = Cm(12.5))

        # section = document.sections[0]
        # footer = section.footer
        # footer.add_paragraph('Reporte Simulacion SIMCOPIAPO | creado el %s' %now.strftime("%d-%m-%Y a las %H:%M"))

        document.add_section(WD_SECTION.NEW_PAGE)

        document.save(self._mf(os.path.join(self.model_report_dir, '1_Evolucion_niveles_y_volumen_embalsado', 'report_out.docx')))

        # Page 4 - Niveles freaticos red monitoreo DGA

        now = dt.datetime.now()

        document = Document(os.path.join(self.model_report_data_dir, '2_Niveles_freaticos_pozos_DGA', 'report_template.docx'))

        style = document.styles['Normal']
        font = style.font
        font.name = 'Avenir Light'
        font.size = Pt(10)

        tables = document.tables

        p = tables[0].rows[0].cells[0].add_paragraph()
        r = p.add_run()
        r.add_picture(os.path.join(self.model_report_data_dir, '2_Niveles_freaticos_pozos_DGA', '1_Baseline', 'niveles_freaticos_pozos_DGA.png'), width = Cm(8.8))

        # ------> REPLACE WITH SCENARIO FIGURE <------
        p = tables[0].rows[0].cells[1].add_paragraph()
        r = p.add_run()
        r.add_picture(self._mf(os.path.join(self.model_report_dir, '2_Niveles_freaticos_pozos_DGA', '2_Scenario', 'niveles_freaticos_pozos_DGA.png')), width = Cm(8.8))

        # section = document.sections[0]
        # footer = section.footer
        # footer.add_paragraph('Reporte Simulacion SIMCOPIAPO | creado el %s' %now.strftime("%d-%m-%Y a las %H:%M"))

        document.add_section(WD_SECTION.NEW_PAGE)

        document.save(self._mf(os.path.join(self.model_report_dir, '2_Niveles_freaticos_pozos_DGA', 'report_out.docx')))

        # Page 5 — Balance acuifero (verano)

        now = dt.datetime.now()

        document = Document(os.path.join(self.model_report_data_dir, '3_Balance_acuifero_verano', 'report_template.docx'))

        style = document.styles['Normal']
        font = style.font
        font.name = 'Avenir Light'
        font.size = Pt(10)

        tables = document.tables

        p = tables[0].rows[0].cells[0].add_paragraph()
        r = p.add_run()
        r.add_picture(os.path.join(self.model_report_data_dir, '3_Balance_acuifero_verano', '1_Baseline', 'balance_acuifero_verano.png'), width = Cm(7.7))

        # ------> REPLACE WITH SCENARIO FIGURE <------
        p = tables[0].rows[0].cells[1].add_paragraph()
        r = p.add_run()
        r.add_picture(self._mf(os.path.join(self.model_report_dir, '3_Balance_acuifero_verano', '2_Scenario', 'balance_acuifero_verano.png')), width = Cm(7.7))

        # section = document.sections[0]
        # footer = section.footer
        # footer.add_paragraph('Reporte Simulacion SIMCOPIAPO | creado el %s' %now.strftime("%d-%m-%Y a las %H:%M"))

        document.add_section(WD_SECTION.NEW_PAGE)

        document.save(self._mf(os.path.join(self.model_report_dir, '3_Balance_acuifero_verano', 'report_out.docx')))

        # Page 6 — Balance acuifero (invierno)

        # ***** 

        # Page 7 — Embalse Lautaro

        now = dt.datetime.now()

        document = Document(os.path.join(self.model_report_data_dir, '5_Embalse_Lautaro', 'report_template.docx'))

        style = document.styles['Normal']
        font = style.font
        font.name = 'Avenir Light'
        font.size = Pt(10)

        tables = document.tables

        p = tables[0].rows[0].cells[0].add_paragraph()
        r = p.add_run()
        r.add_picture(os.path.join(self.model_report_data_dir, '5_Embalse_Lautaro', '1_Baseline', 'embalse_lautaro.png'), width = Cm(20))

        # ------> REPLACE WITH SCENARIO FIGURE <------
        p = tables[0].rows[1].cells[0].add_paragraph()
        r = p.add_run()
        r.add_picture(self._mf(os.path.join(self.model_report_dir, '5_Embalse_Lautaro', '2_Scenario', 'embalse_lautaro.png')), width = Cm(20))

        # section = document.sections[0]
        # footer = section.footer
        # footer.add_paragraph('Reporte Simulacion SIMCOPIAPO | creado el %s' %now.strftime("%d-%m-%Y a las %H:%M"))

        document.add_section(WD_SECTION.NEW_PAGE)

        document.save(self._mf(os.path.join(self.model_report_dir, '5_Embalse_Lautaro', 'report_out.docx')))

        # Page 8 — JVRC

        now = dt.datetime.now()

        document = Document(os.path.join(self.model_report_data_dir, '6_JVRC', 'report_template.docx'))

        style = document.styles['Normal']
        font = style.font
        font.name = 'Avenir Light'
        font.size = Pt(10)

        tables = document.tables

        p = tables[0].rows[0].cells[0].add_paragraph()
        r = p.add_run()
        r.add_picture(os.path.join(self.model_report_data_dir, '6_JVRC', '1_Baseline', 'satisfaccion_dda_ts.png'), width = Cm(9))

        p = tables[0].rows[0].cells[0].add_paragraph()
        r = p.add_run()
        r.add_picture(os.path.join(self.model_report_data_dir, '6_JVRC', '1_Baseline', 'satisfaccion_dda_heatmap.png'), width = Cm(9))

        p = tables[0].rows[0].cells[0].add_paragraph()
        r = p.add_run()
        r.add_picture(os.path.join(self.model_report_data_dir, '6_JVRC', '1_Baseline', 'satisfaccion_dda_radar.png'), width = Cm(5.2))

        # ------> REPLACE WITH SCENARIO FIGURE <------
        p = tables[0].rows[0].cells[1].add_paragraph()
        r = p.add_run()
        r.add_picture(self._mf(os.path.join(self.model_report_dir, '6_JVRC', '2_Scenario', 'satisfaccion_dda_ts.png')), width = Cm(9))

        # ------> REPLACE WITH SCENARIO FIGURE <------
        p = tables[0].rows[0].cells[1].add_paragraph()
        r = p.add_run()
        r.add_picture(self._mf(os.path.join(self.model_report_dir, '6_JVRC', '2_Scenario', 'satisfaccion_dda_heatmap.png')), width = Cm(9))

        # ------> REPLACE WITH SCENARIO FIGURE <------
        p = tables[0].rows[0].cells[1].add_paragraph()
        r = p.add_run()
        r.add_picture(self._mf(os.path.join(self.model_report_dir, '6_JVRC', '2_Scenario', 'satisfaccion_dda_radar.png')), width = Cm(5.2))

        # section = document.sections[0]
        # footer = section.footer
        # footer.add_paragraph('Reporte Simulacion SIMCOPIAPO | creado el %s' %now.strftime("%d-%m-%Y a las %H:%M"))

        document.add_section(WD_SECTION.NEW_PAGE)

        document.save(self._mf(os.path.join(self.model_report_dir, '6_JVRC', 'report_out.docx')))

        # Page 9 — Rio Copiapo

        now = dt.datetime.now()

        document = Document(os.path.join(self.model_report_data_dir, '7_Rio_Copiapo', 'report_template.docx'))

        style = document.styles['Normal']
        font = style.font
        font.name = 'Avenir Light'
        font.size = Pt(10)

        tables = document.tables

        p = tables[0].rows[0].cells[0].add_paragraph()
        r = p.add_run()
        r.add_picture(os.path.join(self.model_report_data_dir, '7_Rio_Copiapo', '1_Baseline', 'rio_copiapo.png'), width = Cm(18.5))

        # ------> REPLACE WITH SCENARIO FIGURE <------
        p = tables[0].rows[1].cells[0].add_paragraph()
        r = p.add_run()
        r.add_picture(self._mf(os.path.join(self.model_report_dir, '7_Rio_Copiapo', '2_Scenario', 'rio_copiapo.png')), width = Cm(18.5))

        # section = document.sections[0]
        # footer = section.footer
        # footer.add_paragraph('Reporte Simulacion SIMCOPIAPO | creado el %s' %now.strftime("%d-%m-%Y a las %H:%M"))

        document.add_section(WD_SECTION.NEW_PAGE)

        document.save(self._mf(os.path.join(self.model_report_dir, '7_Rio_Copiapo', 'report_out.docx')))

        # Page 10 — Caudal Urbano y Humedal

        now = dt.datetime.now()

        document = Document(os.path.join(self.model_report_data_dir, '8_Q_Social_y_Ambiental', 'report_template.docx'))

        style = document.styles['Normal']
        font = style.font
        font.name = 'Avenir Light'
        font.size = Pt(10)

        tables = document.tables

        p = tables[0].rows[0].cells[0].add_paragraph()
        r = p.add_run()
        r.add_picture(os.path.join(self.model_report_data_dir, '8_Q_Social_y_Ambiental', '1_Baseline', 'Q_social_y_ambiental.png'), width = Cm(12.5))

        # ------> REPLACE WITH SCENARIO FIGURE <------
        p = tables[0].rows[0].cells[1].add_paragraph()
        r = p.add_run()
        r.add_picture(self._mf(os.path.join(self.model_report_dir, '8_Q_Social_y_Ambiental', '2_Scenario', 'Q_social_y_ambiental.png')), width = Cm(12.5))

        p = tables[0].rows[1].cells[0].add_paragraph()
        r = p.add_run()
        r.add_picture(os.path.join(self.model_report_data_dir, '8_Q_Social_y_Ambiental', '1_Baseline', 'Q_humedal.png'), width = Cm(12.5))

        # ------> REPLACE WITH SCENARIO FIGURE <------
        p = tables[0].rows[1].cells[1].add_paragraph()
        r = p.add_run()
        r.add_picture(self._mf(os.path.join(self.model_report_dir, '8_Q_Social_y_Ambiental', '2_Scenario', 'Q_humedal.png')), width = Cm(12.5))

        # section = document.sections[0]
        # footer = section.footer
        # footer.add_paragraph('Reporte Simulacion SIMCOPIAPO | creado el %s' %now.strftime("%d-%m-%Y a las %H:%M"))

        document.add_section(WD_SECTION.NEW_PAGE)

        document.save(self._mf(os.path.join(self.model_report_dir, '8_Q_Social_y_Ambiental', 'report_out.docx')))

        # Page 11 — Economicos

        now = dt.datetime.now()

        document = Document(os.path.join(self.model_report_data_dir, '9_Economico', 'report_template.docx'))

        style = document.styles['Normal']
        font = style.font
        font.name = 'Avenir Light'
        font.size = Pt(10)

        tables = document.tables

        p = tables[0].rows[0].cells[0].add_paragraph()
        r = p.add_run()
        r.add_picture(os.path.join(self.model_report_data_dir, '9_Economico', '1_Baseline', 'indicadores_economicos.png'), width = Cm(25.5))

        # ------> REPLACE WITH SCENARIO FIGURE <------
        p = tables[0].rows[1].cells[0].add_paragraph()
        r = p.add_run()
        r.add_picture(self._mf(os.path.join(self.model_report_dir, '9_Economico', '2_Scenario', 'indicadores_economicos.png')), width = Cm(25.5))

        # section = document.sections[0]
        # footer = section.footer
        # footer.add_paragraph('Reporte Simulacion SIMCOPIAPO | creado el %s' %now.strftime("%d-%m-%Y a las %H:%M"))

        document.add_section(WD_SECTION.NEW_PAGE)

        document.save(self._mf(os.path.join(self.model_report_dir, '9_Economico', 'report_out.docx')))

        # Page 12 — Summary radar plot

        # Compile Report

        files = [
            os.path.join(self.model_report_dir, '10_Radar_Resumen', 'report_out.docx'),
            os.path.join(self.model_report_dir, '1_Evolucion_niveles_y_volumen_embalsado', 'report_out.docx'),
            os.path.join(self.model_report_dir, '2_Niveles_freaticos_pozos_DGA', 'report_out.docx'),
            os.path.join(self.model_report_dir, '3_Balance_acuifero_verano', 'report_out.docx'),
            os.path.join(self.model_report_dir, '5_Embalse_Lautaro', 'report_out.docx'),
            os.path.join(self.model_report_dir, '6_JVRC', 'report_out.docx'),
            os.path.join(self.model_report_dir, '7_Rio_Copiapo', 'report_out.docx'),
            os.path.join(self.model_report_dir, '8_Q_Social_y_Ambiental', 'report_out.docx'),
            os.path.join(self.model_report_dir, '9_Economico', 'report_out.docx')
        ]

        def combine_all_docx(filename_master, files_list, filepath_final):
            number_of_sections = len(files_list)
            master = Document_compose(filename_master)
            composer = Composer(master)
            for i in range(0, number_of_sections):
                doc_temp = Document_compose(files_list[i])
                composer.append(doc_temp)
            composer.save(self._mf(filepath_final))

        combine_all_docx(self._mf(os.path.join(self.model_report_dir, '_Tapa_informe', 'cover_out.docx')), files, fp)

    def _mf(self, filepath):
        directory = os.path.dirname(filepath)
        if not os.path.exists(directory):
            os.makedirs(directory)
        return filepath

    def export_results(self):
        if self.report is None:
            raise RuntimeError('Unable to generate report - no model run data available.')
        report_filename = 'sim_copiapo_' + dt.datetime.now().strftime("%d-%m-%Y_%H%M") + '.docx'
        report_filepath = os.path.join(self.model_report_dir, report_filename)
        # Generate report
        self.generate_report(report_filepath)
        # Send report to client
        report_stream = open(report_filepath,'rb')
        report_data = report_stream.read()
        report_data_b64_str = base64.encodebytes(report_data).decode()
        result = {}
        result['filename'] = report_filename
        result['type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        result['data'] = report_data_b64_str
        return result

    def _apply_single_geotiff_layer(self, ml, arr, prop_name, src_srs, src_null_value=-999, dst_null_value=-999):
        # Export as GeoTIFF
        dname = 'geotiff'
        fname = prop_name
        rasterpth = os.path.join(self.working_dir, dname)
        if not os.path.exists(rasterpth):
            os.makedirs(rasterpth)
        if hasattr(arr, 'mask'):
            np.ma.set_fill_value(arr, src_null_value)
            arr = arr.filled()
        ml.sr.export_array(os.path.join(rasterpth, fname + '.tif'), arr)
        # Re-project GeoTIFF
        import gdal
        src_ds = os.path.join(rasterpth, fname + '.tif')
        dst_ds = os.path.join(rasterpth, fname + '_t.tif')
        tmp = gdal.Warp(dst_ds, src_ds, srcSRS=src_srs, dstSRS='EPSG:4326', srcNodata=src_null_value, dstNodata=dst_null_value)
        tmp = None
        self.set_layer(prop_name, sb.LayerGeoTIFFImportedVal(dst_ds))

    def _apply_geotiff_layer(self, ml, array, index, prop_name, src_srs, src_null_value=-999, dst_null_value=-999):
        # Export as GeoTIFF
        dname = 'geotiff'
        now = self._start
        day = now.day
        month = now.month
        year = now.year
        DAYS_IN_MONTH = [31,28,31,30,31,30,31,31,30,31,30,31]
        all_dates = []
        all_values = []
        for t in range(len(array)):
            arr = array[t]
            fname = prop_name + '_t=' + str(t)# * 5)
            rasterpth = os.path.join(self.working_dir, dname)
            if not os.path.exists(rasterpth):
                os.makedirs(rasterpth)
            if hasattr(arr, 'mask'):
                np.ma.set_fill_value(arr, src_null_value)
                arr = arr.filled()
            ml.sr.export_array(os.path.join(rasterpth, fname + '.tif'), arr)
            # Re-project GeoTIFF
            import gdal
            src_ds = os.path.join(rasterpth, fname + '.tif')
            dst_ds = os.path.join(rasterpth, fname + '_t.tif')
            tmp = gdal.Warp(dst_ds, src_ds, srcSRS=src_srs, dstSRS='EPSG:4326', srcNodata=src_null_value, dstNodata=dst_null_value)
            tmp = None
            if t == index:
                # Set layer
                self.set_layer(prop_name, sb.LayerGeoTIFFImportedVal(dst_ds))
            # Prepare layer data
            out_ds = gdal.Open(dst_ds)
            data = out_ds.ReadAsArray()
            out_ds = None
            os.remove(src_ds)
            all_dates.append(now)
            all_values.append(data)
            month = month + 6
            if month > 12:
                month = month - 12
                year = year + 1
            day = DAYS_IN_MONTH[month-1]
            now = dt.datetime(year, month, day)
        # Set layer data
        rows, cols = all_values[0].shape
        # Look at first grid to find valid cells
        valid_cells = []
        for r in range(rows):
            for c in range(cols):
                v = float(all_values[0][r,c])
                has_value = False
                if np.isnan(dst_null_value):
                    if not np.isnan(v):
                        has_value = True
                else:
                    has_value = (v != dst_null_value)
                if has_value:
                    valid_cells.append((r,c))
        # Loop over current cells and set layer data
        for r,c in valid_cells:
            dates = list(range(len(all_dates)))
            values = list(range(len(all_dates)))
            for t in range(len(all_dates)):
                dates[t] = all_dates[t]
                values[t] = float(all_values[t][r,c])
            self.set_layer_data(prop_name, (r,c), sb.TimseriesBaseVal(dates, values))
